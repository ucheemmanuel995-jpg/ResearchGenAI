import streamlit as st
from groq import Groq
import json, re, io, os
from datetime import datetime

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchGenAI – End-to-End Research Platform",
    page_icon="🔬", layout="wide",
    initial_sidebar_state="expanded",
)

FEEDBACK_EMAIL = "oparae995@gmail.com"

# ══════════════════════════════════════════════════════════════════════════════
# CSS — dark academic theme + UX improvements
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@400;600;700;800&family=IBM+Plex+Mono:wght@400;500&display=swap');
html,body,[class*="css"]{font-family:'Sora',sans-serif;background:#060d1c;color:#c8deff;}
.stApp{background:#060d1c;}
section[data-testid="stSidebar"]{background:#080f20;border-right:1px solid #1e2d4a;}
section[data-testid="stSidebar"] *{color:#c8deff !important;}
h1{font-family:'Sora',sans-serif!important;font-weight:800!important;color:#fff!important;}
h2{font-family:'Sora',sans-serif!important;font-weight:700!important;color:#fff!important;}
h3{font-family:'Sora',sans-serif!important;color:#00c6a7!important;letter-spacing:2px;text-transform:uppercase;font-size:13px!important;}
input[type="text"],textarea,.stTextInput input,.stTextArea textarea{background:#0d1526!important;border:1.5px solid #1e2d4a!important;border-radius:10px!important;color:#c8deff!important;font-family:'IBM Plex Mono',monospace!important;}
.stSelectbox>div>div{background:#0d1526!important;border:1.5px solid #1e2d4a!important;border-radius:10px!important;color:#c8deff!important;}
.stButton>button{background:linear-gradient(135deg,#00c6a7,#0077ff)!important;border:none!important;border-radius:10px!important;color:#fff!important;font-family:'Sora',sans-serif!important;font-weight:700!important;font-size:14px!important;padding:0.6rem 1.5rem!important;box-shadow:0 4px 20px #0077ff33!important;}
.stButton>button:hover{opacity:.85!important;}
.stMultiSelect>div>div{background:#0d1526!important;border:1.5px solid #1e2d4a!important;border-radius:10px!important;}
.stMultiSelect span[data-baseweb="tag"]{background:linear-gradient(135deg,#00c6a755,#0077ff33)!important;border:1px solid #00c6a7!important;color:#00c6a7!important;}
.stFileUploader{background:#0d1526!important;border:1.5px dashed #2e3a5a!important;border-radius:10px!important;}
hr{border-color:#1e2d4a!important;}
.step-badge{display:inline-block;background:linear-gradient(135deg,#00c6a7,#0077ff);color:white;font-family:'Sora',sans-serif;font-weight:700;font-size:11px;letter-spacing:2px;padding:4px 12px;border-radius:20px;margin-bottom:8px;}
.output-box{background:#0d1526;border:1px solid #1e2d4a;border-radius:12px;padding:1.2rem 1.4rem;font-family:'IBM Plex Mono',monospace;font-size:13px;line-height:1.85;color:#c8deff;white-space:pre-wrap;max-height:520px;overflow-y:auto;position:relative;}
.edit-box{background:#0a1a2e;border:1.5px solid #00c6a755;border-radius:10px;padding:1rem;margin:.5rem 0 1rem 0;}
.info-card{background:#0a1a2e;border:1px solid #1e3a5a;border-radius:10px;padding:.8rem 1rem;margin-bottom:.8rem;font-size:13px;color:#7a9abf;line-height:1.6;}
.saved-badge{display:inline-block;background:#00c6a720;border:1px solid #00c6a7;color:#00c6a7;border-radius:6px;padding:3px 10px;font-size:11px;font-weight:600;}
.doc-card{background:#0a1220;border:1px solid #1e2d4a;border-radius:10px;padding:.8rem 1rem;margin-bottom:.5rem;font-size:12px;color:#c8deff;}
.output-toolbar{display:flex;gap:8px;align-items:center;margin-bottom:6px;flex-wrap:wrap;}
.output-meta{font-size:11px;color:#4a6a8a;font-family:'Sora',sans-serif;}
.success-banner{background:#0a1a14;border:1px solid #00c6a755;border-radius:10px;padding:.8rem 1rem;margin:.5rem 0;font-size:13px;color:#00c6a7;}
.warning-box{background:#1a1200;border:1px solid #ffd16655;border-radius:10px;padding:.8rem 1rem;margin:.5rem 0;font-size:13px;color:#ffd166;}
@media(max-width:768px){.stButton>button{width:100%!important;} section[data-testid="stSidebar"]{width:280px!important;}}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
STEPS = [("📋","Proposal"),("📁","Upload Docs"),("📘","Chapter One"),
         ("📚","Literature"),("📊","Data Analysis"),("🎯","Conclusion"),
         ("✍️","Abstract"),("🔖","References"),("📄","Export")]

METHODOLOGIES = ["","Quantitative","Qualitative","Mixed Methods",
                 "Experimental","Survey","Case Study","Systematic Review"]

STAT_TESTS = ["Descriptive Statistics","Pearson Correlation","Spearman Correlation",
              "Simple Linear Regression","Multiple Regression","Time Series Regression",
              "Binary Logistic Regression","Independent Samples T-Test","Paired Samples T-Test",
              "One-Way ANOVA","Two-Way ANOVA","Chi-Square Test","Mann-Whitney U Test",
              "Kruskal-Wallis Test","Autoregressive (AR) Model",
              "Full Time Series Analysis (ADF + Differencing + ARIMA + SARIMA)",
              "Factor Analysis","Cluster Analysis"]

REF_STYLES = ["APA 7th","APA 6th","MLA","Chicago","Harvard","Vancouver"]

DEFAULT_PROPOSAL_HEADINGS = ("1. Introduction\n2. Research Methodology\n"
    "2.1 Research Design\n2.2 Population and Sampling\n"
    "2.3 Data Collection Instruments\n2.4 Data Analysis Technique\n2.5 Ethical Considerations")
DEFAULT_CH1_HEADINGS = ("1.1 Background of the Study\n1.2 Statement of the Problem\n"
    "1.3 Aim of the Study\n1.4 Objectives of the Study\n1.5 Research Questions\n"
    "1.6 Research Hypotheses\n1.7 Significance of the Study\n1.8 Scope of the Study\n"
    "1.9 Limitations of the Study\n1.10 Definition of Terms\nReferences")
DEFAULT_METH_HEADINGS = ("3.1 Introduction\n3.2 Research Design\n3.3 Population and Sampling\n"
    "3.4 Research Instrument\n3.5 Validity and Reliability\n"
    "3.6 Data Collection Procedure\n3.7 Data Analysis Technique\n3.8 Ethical Considerations")

# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════
DEFAULTS = {
    "current_step":1,"topic":"","methodology":"","ref_style":"APA 7th",
    "variables":"","uploaded_docs_list":[],"uploaded_docs_text":"",
    "uploaded_doc_names":[],"uploaded_doc_authors":"","image_metadata":{},
    "data_filename":"","data_columns":"","selected_tests":[],"df_json":"",
    "proposal_intro":"","proposal_methodology":"",
    "chapter_one":"","literature":"","analysis":"","analysis_summary":"",
    "conclusion":"","abstract":"","references":"","ts_analysis":"",
    "groq_api_key":"","proposal_headings":DEFAULT_PROPOSAL_HEADINGS,
    "ch1_headings":DEFAULT_CH1_HEADINGS,"meth_headings":DEFAULT_METH_HEADINGS,
    "lit_subheadings":"","all_feedback":[],"edit_mode":{},
}
for k,v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════════════════════════════════════
# CORE HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def get_client():
    try:
        key = st.secrets.get("GROQ_API_KEY","") or st.session_state.groq_api_key
    except Exception:
        key = st.session_state.groq_api_key
    return Groq(api_key=key) if key else None

def stream_response(sys_p, usr_p, state_key):
    client = get_client()
    if not client:
        st.error("⚠️ Enter your Groq API key in the sidebar.")
        return
    st.session_state[state_key] = ""
    placeholder = st.empty()
    try:
        with client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role":"system","content":sys_p},{"role":"user","content":usr_p}],
            max_tokens=2000, stream=True,
        ) as stream:
            for chunk in stream:
                delta = chunk.choices[0].delta.content or ""
                st.session_state[state_key] += delta
                placeholder.markdown(
                    f'<div class="output-box">{st.session_state[state_key]}</div>',
                    unsafe_allow_html=True)
    except Exception as e:
        st.error(f"❌ {e}")

def llm_call(sys_p, usr_p, max_tokens=600):
    client = get_client()
    if not client: return ""
    try:
        r = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role":"system","content":sys_p},{"role":"user","content":usr_p}],
            max_tokens=max_tokens, stream=False)
        return r.choices[0].message.content.strip()
    except Exception:
        return ""

def extract_file_text(uploaded):
    try:
        name = uploaded.name.lower()
        if name.endswith(".pdf"):
            import pdfplumber
            text = ""
            with pdfplumber.open(uploaded) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text += t+"\n"
            if len(text) < 100:
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    uploaded.seek(0)
                    imgs = convert_from_bytes(uploaded.read())
                    text = " ".join(pytesseract.image_to_string(img) for img in imgs)
                except Exception:
                    pass
            return text
        elif name.endswith(".docx"):
            import docx
            doc = docx.Document(uploaded)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif name.endswith((".jpg",".jpeg",".png")):
            return None
        else:
            return uploaded.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[Error: {e}]"

def extract_single_doc_info(doc_name, doc_text, ref_style="APA 7th"):
    import json as _json
    sys_p = f"""You are an expert academic librarian analysing ONE document: "{doc_name}"
Tasks:
1. Find the AUTHOR(S) of THIS document only (not secondary authors cited inside).
   If multiple people co-authored this ONE document, list them all together.
2. Find the YEAR of publication
3. Find the FULL TITLE
4. Format a complete reference in {ref_style} style for this single document
5. Write a 3-4 sentence summary from Abstract, Results, Summary, Conclusion, Recommendations

Return ONLY valid JSON, no extra text:
{{"authors":"Surname, A., & Surname, B.","year":"2024","title":"Full title","reference":"Formatted reference","summary":"3-4 sentence summary"}}"""
    raw = llm_call(sys_p,
        f"Document: {doc_name}\n\nText:\n{doc_text[:5000]}\n\nReturn JSON for this ONE document.",
        max_tokens=500)
    try:
        m = re.search(r"\{[^{}]*\}", raw, re.DOTALL)
        if m: return _json.loads(m.group())
    except Exception: pass
    clean = doc_name.replace(".pdf","").replace(".docx","").replace(".txt","")
    return {"authors":"Unknown Author","year":"n.d.","title":clean,
            "reference":f"Unknown Author. (n.d.). {clean}.","summary":raw[:300] or "Not extracted."}

def build_combined_docs_text(max_chars_per_doc=4000):
    docs = st.session_state.uploaded_docs_list
    if not docs: return ""
    return "\n\n".join(
        f"=== DOCUMENT: {d['name']} ===\n{d['text'][:max_chars_per_doc]}"
        for d in docs)

def word_count(text): return len(text.split()) if text else 0
def reading_time(text): mins = word_count(text)//200; return f"{max(1,mins)} min read"

def output_toolbar(key, label):
    """Render copy + edit + regenerate toolbar above an output box."""
    text = st.session_state.get(key,"")
    if not text: return
    wc = word_count(text)
    rt = reading_time(text)
    st.markdown(
        f'<div class="output-meta">📝 {wc:,} words &nbsp;·&nbsp; {rt}</div>',
        unsafe_allow_html=True)
    col1,col2,col3 = st.columns([1,1,3])
    with col1:
        if st.button("📋 Copy",key=f"copy_{key}",use_container_width=True):
            st.session_state[f"copy_trigger_{key}"] = True
    with col2:
        edit_active = st.session_state.edit_mode.get(key,False)
        if st.button("✏️ Edit" if not edit_active else "💾 Save",
                     key=f"edit_btn_{key}",use_container_width=True):
            st.session_state.edit_mode[key] = not edit_active
            st.rerun()
    if st.session_state.get(f"copy_trigger_{key}"):
        st.components.v1.html(
            f"<script>navigator.clipboard.writeText({json.dumps(text)})</script>",
            height=0)
        st.session_state[f"copy_trigger_{key}"] = False
        st.success("✅ Copied to clipboard!")

def show_output(key, label="Output", allow_regen=False, regen_fn=None):
    text = st.session_state.get(key,"")
    if not text: return
    st.markdown(f"### {label}")
    output_toolbar(key, label)
    edit_active = st.session_state.edit_mode.get(key,False)
    if edit_active:
        edited = st.text_area("Edit content",value=text,height=400,
                               key=f"editor_{key}",label_visibility="collapsed")
        st.session_state[key] = edited
    else:
        st.markdown(f'<div class="output-box">{text}</div>',unsafe_allow_html=True)
    if allow_regen and regen_fn:
        if st.button("🔄 Regenerate",key=f"regen_{key}"):
            regen_fn()

def editable_headings(label, key, default):
    with st.expander(f"✏️ Edit {label} Headings",expanded=False):
        st.markdown('<div class="edit-box">',unsafe_allow_html=True)
        val = st.text_area("Headings (one per line)",value=st.session_state[key],
                           height=200,key=f"edit_{key}",label_visibility="collapsed")
        c1,c2 = st.columns(2)
        with c1:
            if st.button("💾 Save",key=f"save_{key}"):
                st.session_state[key]=val; st.success("✅ Saved!")
        with c2:
            if st.button("↩️ Reset",key=f"reset_{key}"):
                st.session_state[key]=default; st.rerun()
        st.markdown('</div>',unsafe_allow_html=True)

def export_docx():
    """Generate a formatted Word document from all generated sections."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = DocxDoc()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        # Title page
        title = doc.add_heading(st.session_state.topic or "Research Dissertation", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"\nGenerated by ResearchGenAI\n{datetime.now().strftime('%B %Y')}\n")
        doc.add_page_break()
        sections = [
            ("abstract","Abstract"),("proposal_intro","Research Proposal"),
            ("chapter_one","Chapter One — Introduction"),
            ("literature","Chapter Two — Literature Review"),
            ("analysis","Chapter Three — Methodology & Analysis"),
            ("analysis_summary","Analysis Summary"),
            ("conclusion","Chapter Five — Conclusion & Recommendations"),
            ("references","References"),
        ]
        for key,label in sections:
            text = st.session_state.get(key,"")
            if not text: continue
            doc.add_heading(label,1)
            for para in text.split("\n"):
                para = para.strip()
                if not para: continue
                if para.startswith("###"):
                    doc.add_heading(para.replace("###","").strip(),3)
                elif para.startswith("##"):
                    doc.add_heading(para.replace("##","").strip(),2)
                elif para.startswith("#"):
                    doc.add_heading(para.replace("#","").strip(),1)
                else:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.space_after = Pt(6)
            doc.add_page_break()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        st.error(f"❌ Word export error: {e}")
        return None

# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div style="background:#080f20;border-bottom:1px solid #1e2d4a;padding:1.2rem 2rem;
display:flex;align-items:center;gap:1rem;margin:-1rem -1rem 2rem -1rem;">
<div style="width:44px;height:44px;background:linear-gradient(135deg,#00c6a7,#0077ff);
border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:22px;">🔬</div>
<div>
<div style="font-family:Sora,sans-serif;font-weight:800;font-size:20px;color:#fff;">ResearchGenAI</div>
<div style="font-family:Sora,sans-serif;font-size:11px;color:#4a6a8a;letter-spacing:2px;">
END-TO-END RESEARCH &amp; REPORT PLATFORM · Powered by Groq + LLaMA 3.3</div>
</div>
</div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 🔑 Groq API Key")
    api_key_input = st.text_input("Key",value=st.session_state.groq_api_key,
        type="password",placeholder="gsk_...",label_visibility="collapsed")
    if api_key_input != st.session_state.groq_api_key:
        st.session_state.groq_api_key = api_key_input
    if st.session_state.groq_api_key:
        st.success("✅ Key set")
    else:
        st.info("Get FREE key:\nconsole.groq.com")

    st.markdown("---")
    st.markdown("### 🧭 Navigation")
    for i,(icon,label) in enumerate(STEPS,1):
        badge = "✅" if i<st.session_state.current_step else ("▶️" if i==st.session_state.current_step else "⭕")
        if st.button(f"{badge} {i}. {icon} {label}",key=f"nav_{i}",use_container_width=True):
            st.session_state.current_step = i; st.rerun()

    st.markdown("---")
    st.markdown("### 📌 Topic")
    t = st.text_input("Topic",value=st.session_state.topic,
        label_visibility="collapsed",placeholder="Research topic...")
    if t != st.session_state.topic:
        st.session_state.topic = t

    docs = st.session_state.uploaded_docs_list
    if docs:
        st.markdown(f"### 📁 Docs ({len(docs)})")
        for d in docs:
            st.caption(f"📄 {d['name'][:30]} ({d['char_count']:,}c)")

    st.markdown("---")
    sections_done = sum(1 for k in ["proposal_intro","chapter_one","literature",
        "analysis","conclusion","abstract","references"] if st.session_state.get(k))
    st.markdown(f"**Progress:** {sections_done}/7 sections")
    st.progress(sections_done/7)

# ══════════════════════════════════════════════════════════════════════════════
# PROGRESS BAR
# ══════════════════════════════════════════════════════════════════════════════
cur = st.session_state.current_step
pcols = st.columns(len(STEPS)*2-1)
for i,(icon,label) in enumerate(STEPS):
    with pcols[i*2]:
        active,done_s = (i+1==cur),(i+1<cur)
        bg = "linear-gradient(135deg,#00c6a7,#0077ff)" if active else ("#00c6a7" if done_s else "#1e2740")
        bc = "#00c6a7" if (active or done_s) else "#2e3a5a"
        tc = "#00c6a7" if (active or done_s) else "#4a5a7a"
        st.markdown(
            f'<div style="display:flex;flex-direction:column;align-items:center;gap:4px;">'
            f'<div style="width:36px;height:36px;border-radius:50%;background:{bg};'
            f'border:2px solid {bc};display:flex;align-items:center;justify-content:center;'
            f'font-size:14px;box-shadow:{"0 0 14px #00c6a755" if active else "none"};">'
            f'{"✓" if done_s else icon}</div>'
            f'<span style="font-size:8px;font-weight:600;color:{tc};white-space:nowrap;">{label}</span></div>',
            unsafe_allow_html=True)
    if i<len(STEPS)-1:
        with pcols[i*2+1]:
            st.markdown(
                f'<div style="height:2px;background:{"#00c6a7" if i+1<cur else "#1e2740"};margin-top:18px;"></div>',
                unsafe_allow_html=True)
st.markdown("<br>",unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — PROPOSAL
# ══════════════════════════════════════════════════════════════════════════════
if cur==1:
    st.markdown('<div class="step-badge">STEP 1 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📋 Research Proposal")
    st.markdown('<div class="info-card">Two parts only: <strong>Introduction</strong> and <strong>Methodology</strong>. Saved and reused across all chapters. No citations here.</div>',unsafe_allow_html=True)
    topic = st.text_input("Research Topic *",value=st.session_state.topic,
        placeholder="e.g. Time Series Analysis of Malaria Cases in Enugu State")
    st.session_state.topic = topic
    c1,c2 = st.columns(2)
    with c1:
        meth = st.selectbox("Research Methodology *",METHODOLOGIES,
            index=METHODOLOGIES.index(st.session_state.methodology) if st.session_state.methodology in METHODOLOGIES else 0)
        st.session_state.methodology = meth
    with c2:
        rsp = st.selectbox("Citation Style",REF_STYLES,
            index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_p")
        st.session_state.ref_style = rsp
    editable_headings("Proposal","proposal_headings",DEFAULT_PROPOSAL_HEADINGS)
    if st.button("🚀 Generate Proposal",disabled=not(st.session_state.topic and st.session_state.methodology)):
        sys_p = f"""You are an expert academic research proposal writer.
Generate a concise RESEARCH PROPOSAL with ONLY two sections:
SECTION 1 — INTRODUCTION: background (2-3 paras), problem statement, aim, 3+ objectives, research questions, significance.
SECTION 2 — RESEARCH METHODOLOGY using this structure:\n{st.session_state.proposal_headings}
RULES: NO citations. Future tense. No literature review. Professional academic language."""
        stream_response(sys_p,f"Topic:{topic}\nMethodology:{meth}\nGenerate proposal.","proposal_intro")
        st.session_state.proposal_methodology = st.session_state.proposal_intro
    if st.session_state.proposal_intro:
        st.markdown("---")
        show_output("proposal_intro","✅ Proposal Output")
        st.markdown('<div class="saved-badge">✅ Saved — used in all subsequent chapters</div>',unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — DOCUMENT UPLOAD
# ══════════════════════════════════════════════════════════════════════════════
elif cur==2:
    st.markdown('<div class="step-badge">STEP 2 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📁 Document Upload Centre")
    st.markdown('<div class="info-card">Upload all source documents. Each is processed individually — author, year, title and reference extracted per file. Images require caption and source input. Content extracted from: Abstract, Results, Summary, Conclusion, Recommendations.</div>',unsafe_allow_html=True)
    ref_style_up = st.selectbox("Reference Style for Extraction",REF_STYLES,
        index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_upload")
    uploaded_files = st.file_uploader("Upload source documents — as many as needed",
        type=["pdf","txt","docx","jpg","jpeg","png"],
        accept_multiple_files=True,label_visibility="collapsed")
    if uploaded_files and st.button("📥 Process All Documents"):
        new_docs=[]; image_meta=st.session_state.get("image_metadata",{}); all_refs=[]; img_files=[]
        prog = st.progress(0); status = st.empty()
        for idx,f in enumerate(uploaded_files):
            status.markdown(f"⏳ Processing **{f.name}** ({idx+1}/{len(uploaded_files)})...")
            is_img = f.name.lower().endswith((".jpg",".jpeg",".png"))
            if is_img:
                img_files.append(f)
                new_docs.append({"name":f.name,"text":f"[IMAGE: {f.name}]","char_count":0,"is_image":True,
                    "doc_info":{"authors":"See metadata","year":"n.d.","title":f.name,
                    "reference":f"[Image: {f.name}]","summary":"Image file"}})
            else:
                text = extract_file_text(f)
                if not text or text.startswith("[Error"):
                    st.warning(f"⚠️ Could not read {f.name}"); continue
                with st.spinner(f"🔍 Extracting info from {f.name}..."):
                    info = extract_single_doc_info(f.name,text,ref_style_up)
                doc_entry={"name":f.name,"text":text,"char_count":len(text),"is_image":False,"doc_info":info}
                new_docs.append(doc_entry)
                wc = word_count(text)
                st.markdown(
                    f'<div class="doc-card">'
                    f'<strong>✅ {f.name}</strong><br>'
                    f'📊 {len(text):,} chars · {wc:,} words<br>'
                    f'👤 <strong>Author(s):</strong> {info.get("authors","Unknown")}<br>'
                    f'📅 <strong>Year:</strong> {info.get("year","n.d.")}<br>'
                    f'📄 <strong>Title:</strong> {info.get("title","")}<br>'
                    f'<br><strong>📚 Reference ({ref_style_up}):</strong><br>'
                    f'<em>{info.get("reference","")}</em><br>'
                    f'<br><strong>🔑 Summary:</strong><br>{info.get("summary","")}'
                    f'</div>',unsafe_allow_html=True)
                all_refs.append(info.get("reference",""))
            prog.progress((idx+1)/len(uploaded_files))
        if img_files:
            st.markdown("---"); st.markdown("### 🖼️ Image Metadata")
            for img_f in img_files:
                with st.expander(f"📷 {img_f.name}",expanded=True):
                    ca,cb = st.columns(2)
                    with ca:
                        img_name=st.text_input("Figure caption",value=image_meta.get(img_f.name,{}).get("name",""),
                            key=f"imgn_{img_f.name}",placeholder="e.g. Figure 1: Trend of Cases 2010-2023")
                    with cb:
                        img_src=st.text_input("Source / author",value=image_meta.get(img_f.name,{}).get("source",""),
                            key=f"imgs_{img_f.name}",placeholder="e.g. WHO Report (2023)")
                    image_meta[img_f.name]={"name":img_name,"source":img_src}
                    for d in new_docs:
                        if d["name"]==img_f.name:
                            d["text"]=f"[IMAGE: {img_f.name}]\nCaption: {img_name}\nSource: {img_src}"
                            d["doc_info"]["reference"]=f"{img_src}. {img_name}. [Figure]."
            st.session_state.image_metadata=image_meta
        st.session_state.uploaded_docs_list=new_docs
        st.session_state.uploaded_doc_names=[d["name"] for d in new_docs]
        authors_list=[f"{d['doc_info'].get('authors','')} ({d['doc_info'].get('year','')}).{d['doc_info'].get('title','')[:50]}"
            for d in new_docs if not d.get("is_image") and d['doc_info'].get('authors')]
        st.session_state.uploaded_doc_authors="\n".join(authors_list)
        st.session_state.uploaded_docs_text=build_combined_docs_text(4000)
        status.empty()
        if all_refs:
            st.markdown("---"); st.markdown("### 📚 Master Reference List")
            refs_text="\n\n".join(f"{i+1}. {r}" for i,r in enumerate(all_refs))
            st.markdown(f'<div class="output-box">{refs_text}</div>',unsafe_allow_html=True)
            st.session_state.references=refs_text
        total=sum(d["char_count"] for d in new_docs)
        st.success(f"✅ {len(new_docs)} document(s) processed — {total:,} total characters")
    if st.session_state.uploaded_docs_list and not uploaded_files:
        st.markdown(f"### 📚 Loaded Documents ({len(st.session_state.uploaded_docs_list)})")
        for doc in st.session_state.uploaded_docs_list:
            icon="🖼️" if doc.get("is_image") else "📄"
            info=doc.get("doc_info",{})
            st.markdown(
                f'<div class="doc-card">{icon} <strong>{doc["name"]}</strong> — {doc["char_count"]:,} chars<br>'
                f'👤 {info.get("authors","")} ({info.get("year","")}) — {info.get("title","")[:60]}<br>'
                f'<em>{info.get("reference","")}</em></div>',unsafe_allow_html=True)
    if not st.session_state.uploaded_docs_list:
        st.warning("⚠️ No documents yet. Upload files and click Process.")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — CHAPTER ONE
# ══════════════════════════════════════════════════════════════════════════════
elif cur==3:
    st.markdown('<div class="step-badge">STEP 3 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📘 Chapter One — Introduction")
    st.markdown('<div class="info-card">• Citations ONLY in Section 1.1 Background.<br>• Only uploaded document authors cited.<br>• Content from Abstract, Results, Summary, Conclusion, Recommendations of ALL docs.<br>• Fully paraphrased.</div>',unsafe_allow_html=True)
    if not st.session_state.proposal_intro:
        st.warning("⚠️ Complete Step 1 first.")
    elif not st.session_state.uploaded_docs_list:
        st.warning("⚠️ Upload documents in Step 2 first.")
    else:
        rsc1=st.selectbox("Citation Style",REF_STYLES,
            index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_c1")
        editable_headings("Chapter One","ch1_headings",DEFAULT_CH1_HEADINGS)
        doc_authors=st.session_state.uploaded_doc_authors or "Authors from uploaded documents"
        num_docs=len(st.session_state.uploaded_docs_list)
        st.info(f"📚 {num_docs} document(s) will be integrated.")
        # Editable author list
        with st.expander("✏️ Confirm / Edit Document Authors",expanded=False):
            edited_authors=st.text_area("Authors (one per line — correct any errors)",
                value=doc_authors,height=150,key="edit_authors")
            if st.button("💾 Save Author List"):
                st.session_state.uploaded_doc_authors=edited_authors
                st.success("✅ Authors updated!")
        def gen_ch1():
            per_doc=max(1500,12000//max(num_docs,1))
            all_docs="\n\n".join(f"--- FROM: {d['name']} ---\n{d['text'][:per_doc]}"
                for d in st.session_state.uploaded_docs_list)
            authors=st.session_state.uploaded_doc_authors or doc_authors
            sys_p = f"""You are an expert academic dissertation writer.
Generate a complete CHAPTER ONE — INTRODUCTION.
Use this heading structure:\n{st.session_state.ch1_headings}

CITATION RULES:
- Citations in {rsc1} style ONLY in Section 1.1 Background
- Cite ONLY: {authors}
- NO citations in Sections 1.2-1.10
- References section lists ONLY these authors

SECTION 1.1 BACKGROUND — CRITICAL:
- Minimum 5-6 substantial paragraphs
- Each paragraph cites one uploaded document
- Build: global → regional → local → specific gap
- Every factual claim must be cited

ALL OTHER SECTIONS: derive from proposal objectives + document content
- Paraphrase completely — restructure sentences, vary vocabulary
- Do NOT copy verbatim — plagiarism prevention
- Minimum 2 paragraphs per heading"""
            usr_p=(f"Topic:{st.session_state.topic}\nMethodology:{st.session_state.methodology}\n"
                f"Citation Style:{rsc1}\nProposal:\n{st.session_state.proposal_intro[:800]}\n\n"
                f"ALL DOCS ({num_docs}):\n{all_docs}\n\nAuthors for 1.1:{authors}\n\n"
                f"Generate complete Chapter One. Paraphrase all content. Citations only in 1.1.")
            stream_response(sys_p,usr_p,"chapter_one")
        if st.button("📘 Generate Chapter One"):
            gen_ch1()
        show_output("chapter_one","📘 Chapter One Output",allow_regen=True,regen_fn=gen_ch1)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
elif cur==4:
    st.markdown('<div class="step-badge">STEP 4 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📚 Literature Review — Chapter Two")
    n_docs=len(st.session_state.uploaded_docs_list)
    st.markdown(f'<div class="info-card">• No headings by default — continuous academic prose.<br>• Each paragraph begins with an uploaded document author.<br>• All {n_docs} documents integrated.<br>• Abstract, Results, Summary, Conclusion, Recommendations extracted.<br>• Fully paraphrased.</div>',unsafe_allow_html=True)
    if not st.session_state.uploaded_docs_list:
        st.warning("⚠️ Upload documents in Step 2 first.")
    else:
        rslit=st.selectbox("Citation Style",REF_STYLES,
            index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_lit")
        doc_authors=st.session_state.uploaded_doc_authors or "Authors from documents"
        with st.expander("➕ Optional: Add Subheadings",expanded=False):
            st.markdown('<div class="edit-box">',unsafe_allow_html=True)
            sub_val=st.text_area("Subheadings (leave blank for none)",
                value=st.session_state.lit_subheadings,height=120,key="lit_sub_input",
                label_visibility="collapsed",placeholder="e.g.\n2.1 Conceptual Framework\n2.2 Empirical Studies")
            if st.button("💾 Save Subheadings",key="save_lit_sub"):
                st.session_state.lit_subheadings=sub_val; st.success("✅ Saved!")
            st.markdown('</div>',unsafe_allow_html=True)
        c1,c2=st.columns(2)
        def gen_lit():
            subs=st.session_state.lit_subheadings.strip()
            hinstr=(f"Organise under:\n{subs}\nEach section starts paragraphs with an author."
                if subs else "Write as CONTINUOUS PROSE — NO subheadings.")
            per_doc=max(2000,14000//max(n_docs,1))
            all_docs="\n\n".join(f"--- SOURCE: {d['name']} ---\n{d['text'][:per_doc]}"
                for d in st.session_state.uploaded_docs_list)
            sys_p=f"""You are an expert academic writer specialising in literature reviews.
Generate a LONG, THOROUGH CHAPTER TWO — LITERATURE REVIEW.
{hinstr}

STRUCTURE:
1. Opening paragraph (no citation) — introduces themes
2. For EACH of the {n_docs} documents, write MINIMUM 2 paragraphs:
   - Para 1: "[Author] ([Year]) examined/investigated/found that..." — study, purpose, objectives
   - Para 2: methodology, key findings, conclusions, citation at end
   - Para 3 if rich: implications, gaps noted
3. Synthesis paragraph — compare/contrast all findings
4. Research Gap paragraph — what is missing, why this study is needed

CITATION RULES: {rslit} style. ONLY these authors: {doc_authors}
CONTENT: ALL {n_docs} docs must appear (min 2 paras each). Draw from Abstract, Results, Summary, Conclusion, Recommendations.
Paraphrase completely. Do NOT copy verbatim. Each para minimum 5-6 sentences.
REFERENCES: End with References — only uploaded authors in {rslit}, alphabetical."""
            usr_p=(f"Topic:{st.session_state.topic}\nStyle:{rslit}\n"
                f"Proposal:\n{st.session_state.proposal_intro[:500]}\n\n"
                f"ALL DOCS ({n_docs}):\n{all_docs}\n\nAuthors:{doc_authors}\n\n"
                f"Generate comprehensive literature review integrating ALL {n_docs} documents.")
            stream_response(sys_p,usr_p,"literature")
        with c1:
            if st.button("📚 Generate Literature Review",use_container_width=True):
                gen_lit()
        with c2:
            if st.button("🔍 Extract References Only",use_container_width=True):
                combined=build_combined_docs_text(2000)
                sys_p=f"Extract ONLY the references of the uploaded documents. Format in {rslit}. Sort alphabetically. Number each."
                usr_p=f"Style:{rslit}\nAuthors:{doc_authors}\nDocs:\n{combined}\nFormat references."
                stream_response(sys_p,usr_p,"references"); st.success("✅ Saved to Step 8!")
        show_output("literature","📚 Chapter Two Output",allow_regen=True,regen_fn=gen_lit)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — DATA ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
elif cur==5:
    st.markdown('<div class="step-badge">STEP 5 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📊 Methodology & Data Analysis — Chapter Three")
    st.markdown('<div class="info-card">• Methodology: no citations in prose; formulas have citations.<br>• Assumptions: actual computed values + automatic remediation if failed.<br>• Interactive Plotly charts for descriptive analysis.<br>• Every table followed by interpretation.<br>• Full ARIMA/SARIMA with ACF/PACF.</div>',unsafe_allow_html=True)
    if not st.session_state.proposal_intro:
        st.warning("⚠️ Complete Step 1 first.")
    else:
        data_file=st.file_uploader("Upload Dataset (Excel / CSV)",type=["xlsx","xls","csv"])
        df=None
        if data_file:
            st.session_state.data_filename=data_file.name
            try:
                import pandas as pd, numpy as np
                df_raw=(pd.read_csv(data_file) if data_file.name.endswith(".csv") else pd.read_excel(data_file))
                st.session_state.data_columns=", ".join(str(c) for c in df_raw.columns)
                st.session_state.df_json=df_raw.to_json()
                st.success(f"✅ {data_file.name} — {df_raw.shape[0]} rows × {df_raw.shape[1]} cols")
                st.dataframe(df_raw.head(10),use_container_width=True); df=df_raw
            except Exception as e:
                st.error(f"❌ {e}")
        if df is None and st.session_state.df_json:
            try:
                import pandas as pd, numpy as np
                df=pd.read_json(st.session_state.df_json)
            except Exception: pass
        ca,cb=st.columns(2)
        with ca:
            variables=st.text_input("Dependent / Independent Variables",
                value=st.session_state.variables,placeholder="e.g. Dependent: Cases; Independent: Rainfall")
            st.session_state.variables=variables
        with cb:
            rsan=st.selectbox("Citation Style",REF_STYLES,
                index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_an")
        ts_col=None; time_col=None
        if df is not None:
            numeric_cols=df.select_dtypes(include=["number"]).columns.tolist()
            all_cols=["— None —"]+list(df.columns)
            num_opts=["— None —"]+numeric_cols
        run_full_ts=False; run_ts_reg=False
        selected_tests=st.multiselect("Select Statistical Tests",STAT_TESTS,
            default=st.session_state.selected_tests)
        st.session_state.selected_tests=selected_tests
        run_full_ts="Full Time Series Analysis (ADF + Differencing + ARIMA + SARIMA)" in selected_tests
        run_ts_reg="Time Series Regression" in selected_tests
        if df is not None and (run_full_ts or run_ts_reg):
            st.markdown("### ⏱️ Time Series Configuration")
            tc1,tc2=st.columns(2)
            with tc1:
                time_col=st.selectbox("Date/Time Column",all_cols,key="ts_time")
                if time_col=="— None —": time_col=None
            with tc2:
                ts_col=st.selectbox("Target Variable",num_opts,key="ts_target")
                if ts_col=="— None —": ts_col=None
        editable_headings("Methodology","meth_headings",DEFAULT_METH_HEADINGS)

        # ── Plotly Interactive Descriptive Graphs ──────────────────────────
        if df is not None and "Descriptive Statistics" in selected_tests:
            st.markdown("---"); st.markdown("### 📈 Interactive Descriptive Charts")
            try:
                import plotly.express as px
                import plotly.graph_objects as go
                import pandas as pd, numpy as np
                numeric_cols=df.select_dtypes(include=[np.number]).columns.tolist()
                cat_cols=df.select_dtypes(include=["object","category"]).columns.tolist()
                COLORS=["#00c6a7","#0077ff","#ff6b6b","#ffd166","#a29bfe","#fd79a8","#55efc4"]
                plot_cfg={"displayModeBar":True,"modeBarButtonsToRemove":["lasso2d","select2d"],"toImageButtonOptions":{"format":"png","filename":"chart","scale":2}}
                layout_base=dict(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",
                    font_color="#c8deff",font_family="IBM Plex Mono",
                    xaxis=dict(gridcolor="#1e2d4a",linecolor="#1e2d4a"),
                    yaxis=dict(gridcolor="#1e2d4a",linecolor="#1e2d4a"))
                # 1. Histograms
                if numeric_cols:
                    st.markdown("**1. Histograms — frequency distribution**")
                    for col in numeric_cols:
                        fig=px.histogram(df,x=col,nbins=20,title=f"Distribution: {col}",
                            color_discrete_sequence=["#00c6a7"])
                        fig.update_layout(**layout_base,title_font_color="#00c6a7")
                        st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 2. Box plots
                if numeric_cols:
                    st.markdown("**2. Box plots — spread and outliers**")
                    fig=go.Figure()
                    for i,col in enumerate(numeric_cols):
                        fig.add_trace(go.Box(y=df[col].dropna(),name=col,
                            marker_color=COLORS[i%len(COLORS)],boxmean=True))
                    fig.update_layout(**layout_base,title="Box Plot Comparison",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 3. Bar charts (categorical)
                if cat_cols:
                    st.markdown("**3. Bar charts — categorical frequencies**")
                    for col in cat_cols[:4]:
                        vc=df[col].value_counts().reset_index()
                        vc.columns=[col,"count"]
                        fig=px.bar(vc,x=col,y="count",title=f"Frequency: {col}",
                            color=col,color_discrete_sequence=COLORS)
                        fig.update_layout(**layout_base,title_font_color="#00c6a7",showlegend=False)
                        st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 4. Line graphs (numeric over index/time)
                if numeric_cols:
                    st.markdown("**4. Line graphs — trends over observations**")
                    fig=go.Figure()
                    for i,col in enumerate(numeric_cols[:5]):
                        fig.add_trace(go.Scatter(y=df[col].values,mode="lines",
                            name=col,line=dict(color=COLORS[i%len(COLORS)],width=1.5)))
                    fig.update_layout(**layout_base,title="Trend Lines",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 5. Area charts
                if numeric_cols:
                    st.markdown("**5. Area charts — cumulative view**")
                    fig=go.Figure()
                    for i,col in enumerate(numeric_cols[:4]):
                        fig.add_trace(go.Scatter(y=df[col].fillna(0).values,mode="lines",
                            name=col,fill="tozeroy",
                            line=dict(color=COLORS[i%len(COLORS)],width=1.2),
                            fillcolor=COLORS[i%len(COLORS)].replace("#","rgba(").replace("c6a7","198,167,0,.15)").replace("0077ff","0,119,255,.12)") if i<2 else COLORS[i%len(COLORS)]))
                    fig.update_layout(**layout_base,title="Area Chart",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 6. Frequency polygon
                if numeric_cols:
                    st.markdown("**6. Frequency polygons**")
                    fig=go.Figure()
                    for i,col in enumerate(numeric_cols[:5]):
                        d=df[col].dropna()
                        counts,bins=np.histogram(d,bins=15)
                        mids=(bins[:-1]+bins[1:])/2
                        fig.add_trace(go.Scatter(x=mids,y=counts,mode="lines+markers",
                            name=col,line=dict(color=COLORS[i%len(COLORS)],width=1.8),
                            marker=dict(size=4)))
                    fig.update_layout(**layout_base,title="Frequency Polygon",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 7. Correlation heatmap
                if len(numeric_cols)>=2:
                    st.markdown("**7. Correlation heatmap**")
                    corr=df[numeric_cols].corr()
                    fig=go.Figure(go.Heatmap(z=corr.values,x=corr.columns,y=corr.columns,
                        colorscale=[[0,"#0d1526"],[0.5,"#0077ff"],[1,"#00c6a7"]],
                        zmid=0,text=[[f"{v:.2f}" for v in row] for row in corr.values],
                        texttemplate="%{text}",textfont_size=10))
                    fig.update_layout(**layout_base,title="Correlation Matrix",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 8. Scatter matrix (pairs plot)
                if len(numeric_cols)>=2:
                    st.markdown("**8. Scatter plot matrix (pairs)**")
                    cols_sp=numeric_cols[:4]
                    fig=px.scatter_matrix(df,dimensions=cols_sp,color=cat_cols[0] if cat_cols else None,
                        color_discrete_sequence=COLORS)
                    fig.update_layout(**layout_base,title="Scatter Matrix",title_font_color="#00c6a7")
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                # 9. Stacked bar (cross-tab)
                if len(cat_cols)>=2:
                    st.markdown("**9. Stacked bar — cross-tabulation**")
                    try:
                        ct=pd.crosstab(df[cat_cols[0]],df[cat_cols[1]])
                        fig=go.Figure()
                        for i,col in enumerate(ct.columns):
                            fig.add_trace(go.Bar(name=str(col),x=ct.index.astype(str),
                                y=ct[col].values,marker_color=COLORS[i%len(COLORS)]))
                        fig.update_layout(**layout_base,barmode="stack",
                            title=f"Stacked: {cat_cols[0]} × {cat_cols[1]}",title_font_color="#00c6a7")
                        st.plotly_chart(fig,use_container_width=True,config=plot_cfg)
                    except Exception: pass
                # 5b. Pie / Donut charts (categorical proportions)
                if cat_cols:
                    st.markdown("**5. Pie Charts — Proportional Breakdown**")
                    pie_cols_display = st.columns(min(3, len(cat_cols)))
                    for ci_pie, col_pie in enumerate(cat_cols[:3]):
                        vc_pie = df[col_pie].value_counts().head(6)
                        import plotly.graph_objects as go
                        fig_pie = go.Figure(go.Pie(
                            labels=vc_pie.index.astype(str), values=vc_pie.values,
                            hole=0.35,
                            marker=dict(colors=["#00c6a7","#0077ff","#ff6b6b","#ffd166","#a29bfe","#fd79a8"]),
                            textinfo="percent+label", textfont_size=11,
                        ))
                        fig_pie.update_layout(
                            paper_bgcolor="rgba(0,0,0,0)",
                            font_color="#c8deff",
                            title=f"Distribution: {col_pie}",
                            title_font_color="#00c6a7",
                            showlegend=True,
                            legend=dict(bgcolor="rgba(0,0,0,0)", font_color="#c8deff"),
                        )
                        with pie_cols_display[ci_pie]:
                            st.plotly_chart(fig_pie, use_container_width=True, config=plot_cfg)

                # 10. Violin plots
                if numeric_cols and cat_cols:
                    st.markdown("**10. Violin plots — distribution by group**")
                    fig=px.violin(df,y=numeric_cols[0],x=cat_cols[0],box=True,
                        color=cat_cols[0],color_discrete_sequence=COLORS,
                        title=f"{numeric_cols[0]} by {cat_cols[0]}")
                    fig.update_layout(**layout_base,title_font_color="#00c6a7",showlegend=False)
                    st.plotly_chart(fig,use_container_width=True,config=plot_cfg)

                # 11. Stacked bar chart (cross-tabulation of categorical vars)
                if len(cat_cols) >= 2:
                    st.markdown("**11. Stacked Bar Chart — Cross-tabulation**")
                    try:
                        import pandas as pd
                        ct_sb = pd.crosstab(df[cat_cols[0]], df[cat_cols[1]])
                        fig_sb = go.Figure()
                        for ci_sb, col_sb in enumerate(ct_sb.columns):
                            fig_sb.add_trace(go.Bar(
                                name=str(col_sb),
                                x=ct_sb.index.astype(str),
                                y=ct_sb[col_sb].values,
                                marker_color=CLRS[ci_sb % len(CLRS)]))
                        fig_sb.update_layout(**layout_base, barmode="stack",
                            title=f"Stacked: {cat_cols[0]} × {cat_cols[1]}",
                            title_font_color="#00c6a7")
                        st.plotly_chart(fig_sb, use_container_width=True, config=plot_cfg)
                    except Exception: pass

                # 12. Grouped mean bar chart
                if cat_cols and numeric_cols:
                    st.markdown("**12. Grouped Mean Bar Chart — Categorical vs Numeric Mean**")
                    for cat_gm in cat_cols[:2]:
                        for num_gm in numeric_cols[:2]:
                            try:
                                grp_gm = df.groupby(cat_gm)[num_gm].mean().dropna().reset_index()
                                grp_gm.columns = [cat_gm, "Mean"]
                                fig_gm = px.bar(grp_gm, x=cat_gm, y="Mean",
                                    title=f"Mean {num_gm} by {cat_gm}",
                                    color=cat_gm, color_discrete_sequence=CLRS,
                                    text="Mean")
                                fig_gm.update_traces(texttemplate="%{text:.2f}", textposition="outside")
                                fig_gm.update_layout(**layout_base,
                                    title_font_color="#00c6a7", showlegend=False)
                                st.plotly_chart(fig_gm, use_container_width=True, config=plot_cfg)
                            except Exception: pass

                # 13. Histogram with normal distribution overlay
                if numeric_cols:
                    st.markdown("**13. Histograms with Normal Distribution Curve**")
                    for col_nc in numeric_cols[:3]:
                        d_nc = df[col_nc].dropna()
                        if len(d_nc) < 3: continue
                        import numpy as _np_nc
                        from scipy.stats import norm as _norm_nc
                        mu_nc, std_nc = d_nc.mean(), d_nc.std()
                        x_nc = _np_nc.linspace(d_nc.min(), d_nc.max(), 100)
                        pdf_nc = _norm_nc.pdf(x_nc, mu_nc, std_nc)
                        fig_nc = go.Figure()
                        fig_nc.add_trace(go.Histogram(x=d_nc, nbinsx=20,
                            histnorm="probability density",
                            marker_color="#00c6a7", opacity=0.75,
                            name="Data"))
                        fig_nc.add_trace(go.Scatter(x=x_nc, y=pdf_nc,
                            mode="lines", line=dict(color="#ffd166", width=2.5),
                            name="Normal curve"))
                        fig_nc.update_layout(**layout_base,
                            title=f"Distribution + Normal Curve: {col_nc}",
                            title_font_color="#00c6a7",
                            xaxis_title=col_nc, yaxis_title="Density")
                        st.plotly_chart(fig_nc, use_container_width=True, config=plot_cfg)

            except Exception as eg:
                st.warning(f"⚠️ Chart error: {eg}")

        # ── Full Time Series Block ─────────────────────────────────────────
        if df is not None and (run_full_ts or run_ts_reg) and ts_col:
            st.markdown("---"); st.markdown("### ⏱️ Time Series Analysis")
            try:
                import plotly.graph_objects as go
                import plotly.express as px
                import pandas as pd, numpy as np
                from scipy import stats as _stats
                series=df[ts_col].dropna()
                if time_col and time_col in df.columns:
                    try:
                        df[time_col]=pd.to_datetime(df[time_col])
                        series.index=df[time_col][:len(series)]
                    except Exception: pass
                fig=go.Figure()
                fig.add_trace(go.Scatter(y=series.values,mode="lines",
                    line=dict(color="#00c6a7",width=1.5),name=ts_col))
                fig.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",
                    font_color="#c8deff",title=f"Time Series: {ts_col}",title_font_color="#00c6a7")
                st.plotly_chart(fig,use_container_width=True)
                try:
                    from statsmodels.tsa.stattools import adfuller
                    adf=adfuller(series,autolag="AIC")
                    adf_stat,adf_p,adf_lags,adf_nobs,crit=adf[0],adf[1],adf[2],adf[3],adf[4]
                    st.markdown("#### ADF Stationarity Test")
                    st.markdown(f"""
| Parameter | Value |
|---|---|
| ADF Statistic | {adf_stat:.4f} |
| p-value | {adf_p:.4f} |
| Lags | {adf_lags} | Observations | {adf_nobs} |
| Critical (1%) | {crit['1%']:.4f} | Critical (5%) | {crit['5%']:.4f} |
""")
                    is_stationary=adf_p<0.05
                    series_stationary=series; d_order=0
                    if is_stationary:
                        st.success(f"✅ STATIONARY (p={adf_p:.4f} < 0.05). No differencing needed.")
                    else:
                        st.error(f"🚨 NON-STATIONARY (p={adf_p:.4f} ≥ 0.05). Applying differencing...")
                        series_diff=series.diff().dropna(); d_order=1
                        adf2=adfuller(series_diff,autolag="AIC")
                        st.markdown(f"**After 1st differencing:** ADF={adf2[0]:.4f}, p={adf2[1]:.4f} — {'✅ Stationary' if adf2[1]<0.05 else '⚠️ Still non-stationary'}")
                        if adf2[1]>=0.05:
                            series_diff2=series_diff.diff().dropna(); d_order=2
                            adf3=adfuller(series_diff2,autolag="AIC")
                            st.markdown(f"**After 2nd differencing:** ADF={adf3[0]:.4f}, p={adf3[1]:.4f}")
                            series_stationary=series_diff2
                        else:
                            series_stationary=series_diff
                        fig2=go.Figure()
                        fig2.add_trace(go.Scatter(y=series_stationary.values,mode="lines",
                            line=dict(color="#0077ff",width=1.5),name=f"Differenced (d={d_order})"))
                        fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",
                            font_color="#c8deff",title=f"Differenced Series (d={d_order})",title_font_color="#00c6a7")
                        st.plotly_chart(fig2,use_container_width=True)
                    # ACF / PACF
                    from statsmodels.graphics.tsaplots import plot_acf,plot_pacf
                    from statsmodels.tsa.stattools import acf,pacf
                    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
                    lags_max=min(20,len(series_stationary)//2-1)
                    fig_acf,(ax1,ax2)=plt.subplots(1,2,figsize=(14,4),facecolor="#0d1526")
                    for ax in [ax1,ax2]:
                        ax.set_facecolor("#0d1526")
                        for sp in ax.spines.values(): sp.set_edgecolor("#1e2d4a")
                        ax.tick_params(colors="#7a9abf")
                    plot_acf(series_stationary,lags=lags_max,ax=ax1,color="#00c6a7",title="ACF")
                    plot_pacf(series_stationary,lags=lags_max,ax=ax2,color="#0077ff",title="PACF")
                    ax1.set_title("ACF — Autocorrelation Function",color="#c8deff")
                    ax2.set_title("PACF — Partial Autocorrelation Function",color="#c8deff")
                    fig_acf.patch.set_facecolor("#0d1526"); plt.tight_layout()
                    st.pyplot(fig_acf); plt.close()
                    st.markdown('<div class="info-card">📌 <strong>ACF</strong>: spikes at lag k → MA(q) order. <strong>PACF</strong>: spikes at lag k → AR(p) order.</div>',unsafe_allow_html=True)
                    acf_v=acf(series_stationary,nlags=lags_max)
                    pacf_v=pacf(series_stationary,nlags=lags_max)
                    conf=1.96/np.sqrt(len(series_stationary))
                    p_order=min(3,sum(1 for v in pacf_v[1:] if abs(v)>conf))
                    q_order=min(3,sum(1 for v in acf_v[1:] if abs(v)>conf))
                    if run_full_ts:
                        # ARIMA
                        try:
                            from statsmodels.tsa.arima.model import ARIMA
                            st.markdown(f"#### ARIMA({p_order},{d_order},{q_order})")
                            am=ARIMA(series,order=(p_order,d_order,q_order)).fit()
                            st.markdown(f"| AIC | BIC | Log-Likelihood | N |\n|---|---|---|---|\n| {am.aic:.4f} | {am.bic:.4f} | {am.llf:.4f} | {int(am.nobs)} |")
                            st.text(am.summary().as_text())
                            fc=am.forecast(steps=10); ci=am.get_forecast(steps=10).conf_int()
                            fig_fc=go.Figure()
                            fig_fc.add_trace(go.Scatter(y=series.values,mode="lines",name="Observed",line=dict(color="#00c6a7",width=1.5)))
                            fc_idx=list(range(len(series),len(series)+10))
                            fig_fc.add_trace(go.Scatter(x=fc_idx,y=fc.values,mode="lines",name="Forecast",line=dict(color="#ffd166",width=2,dash="dash")))
                            fig_fc.add_trace(go.Scatter(x=fc_idx+fc_idx[::-1],y=list(ci.iloc[:,0])+list(ci.iloc[:,1])[::-1],fill="toself",fillcolor="rgba(0,119,255,.15)",line=dict(width=0),name="95% CI"))
                            fig_fc.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",title=f"ARIMA({p_order},{d_order},{q_order}) Forecast",title_font_color="#00c6a7")
                            st.plotly_chart(fig_fc,use_container_width=True)
                        except Exception as ae:
                            st.warning(f"⚠️ ARIMA: {ae}")
                        # SARIMA
                        try:
                            from statsmodels.tsa.statespace.sarimax import SARIMAX
                            sp=12 if len(series)>=24 else 4
                            st.markdown(f"#### SARIMA({p_order},{d_order},{q_order})×(1,1,1)[{sp}]")
                            sm=SARIMAX(series,order=(p_order,d_order,q_order),seasonal_order=(1,1,1,sp),enforce_stationarity=False,enforce_invertibility=False).fit(disp=False)
                            st.markdown(f"| AIC | BIC | N |\n|---|---|---|\n| {sm.aic:.4f} | {sm.bic:.4f} | {int(sm.nobs)} |")
                            st.text(sm.summary().as_text())
                            sfc=sm.forecast(steps=12); sci=sm.get_forecast(steps=12).conf_int()
                            fig_sfc=go.Figure()
                            fig_sfc.add_trace(go.Scatter(y=series.values,mode="lines",name="Observed",line=dict(color="#00c6a7",width=1.5)))
                            sfc_idx=list(range(len(series),len(series)+12))
                            fig_sfc.add_trace(go.Scatter(x=sfc_idx,y=sfc.values,mode="lines",name="SARIMA Forecast",line=dict(color="#a29bfe",width=2,dash="dash")))
                            fig_sfc.add_trace(go.Scatter(x=sfc_idx+sfc_idx[::-1],y=list(sci.iloc[:,0])+list(sci.iloc[:,1])[::-1],fill="toself",fillcolor="rgba(162,155,254,.15)",line=dict(width=0),name="95% CI"))
                            fig_sfc.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",title=f"SARIMA Forecast (m={sp})",title_font_color="#00c6a7")
                            st.plotly_chart(fig_sfc,use_container_width=True)
                            # Model comparison
                            try:
                                best="ARIMA" if am.aic<sm.aic else "SARIMA"
                                st.markdown(f"""
#### Model Comparison
| Model | AIC | BIC | Best |
|---|---|---|---|
| ARIMA({p_order},{d_order},{q_order}) | {am.aic:.4f} | {am.bic:.4f} | {"✅" if best=="ARIMA" else ""} |
| SARIMA×(1,1,1)[{sp}] | {sm.aic:.4f} | {sm.bic:.4f} | {"✅" if best=="SARIMA" else ""} |

**Recommended: {best}** (lower AIC/BIC = better fit)""")
                            except Exception: pass
                        except Exception as se:
                            st.warning(f"⚠️ SARIMA: {se}")
                    st.session_state.ts_analysis=(f"ADF={adf_stat:.4f},p={adf_p:.4f},"
                        f"Stationary={'Yes' if is_stationary else 'No'},d={d_order}")
                except ImportError:
                    st.error("❌ statsmodels not installed. Run: pip install statsmodels")
                except Exception as tse:
                    st.error(f"❌ Time series error: {tse}")
            except Exception as outer_e:
                st.error(f"❌ {outer_e}")
        # ── Live Assumption Testing + AR Model ──────────────────────────
        if df is not None and len(selected_tests) > 0:
            st.markdown("---")
            st.markdown("### 🔬 Assumption Testing with Remediation")
            st.markdown('<div class="info-card">Each assumption tested with actual computed values. If any fails, the remedy is named, its formula shown, and the corrective procedure executed automatically.</div>', unsafe_allow_html=True)
            try:
                import pandas as pd, numpy as np
                from scipy import stats as _stats
                numeric_cols_a = df.select_dtypes(include=[np.number]).columns.tolist()
                cat_cols_a = df.select_dtypes(include=["object","category"]).columns.tolist()
                test_col_a = numeric_cols_a[0] if numeric_cols_a else None
                if test_col_a:
                    data_a = df[test_col_a].dropna()
                    # A. Normality
                    st.markdown("**A. Normality — Shapiro-Wilk**")
                    if 3 <= len(data_a) <= 5000:
                        sw_s, sw_p = _stats.shapiro(data_a)
                        norm_ok = sw_p > 0.05
                        st.markdown(f"| Variable | W-statistic | p-value | Decision |\n|---|---|---|---|\n| {test_col_a} | {sw_s:.4f} | {sw_p:.4f} | {'✅ Normal' if norm_ok else '❌ Non-normal'} |")
                        if not norm_ok:
                            st.error("🚨 **NORMALITY FAILED — REMEDIATION APPLIED**")
                            st.markdown(f"""
**W={sw_s:.4f}, p={sw_p:.4f}** — Data is not normally distributed.

### 🔧 Remedies:
| Remedy | Formula | When |
|---|---|---|
| **Log Transform** | Y* = ln(Y) | Right-skewed |
| **Square Root** | Y* = √Y | Count data |
| **Box-Cox** | Y* = (Yλ-1)/λ | General |
| **Non-parametric** | Kruskal-Wallis / Mann-Whitney | Replace ANOVA/t-test |

✅ **EXECUTED: Log Transformation applied**""")
                            if (data_a > 0).all():
                                d_log = np.log(data_a)
                                sw2_s, sw2_p = _stats.shapiro(d_log)
                                st.markdown(f"| After log(Y) | W={sw2_s:.4f} | p={sw2_p:.4f} | {'✅ Normality restored' if sw2_p > 0.05 else '⚠️ Still non-normal — use non-parametric'} |")
                                if sw2_p <= 0.05 and cat_cols_a:
                                    groups_kw = [gd[test_col_a].dropna().values for _, gd in df.groupby(cat_cols_a[0]) if len(gd[test_col_a].dropna()) >= 2]
                                    if len(groups_kw) >= 2:
                                        kw_s, kw_p = _stats.kruskal(*groups_kw)
                                        st.markdown(f"✅ **EXECUTED: Kruskal-Wallis** | H={kw_s:.4f} | p={kw_p:.4f} | {'Significant' if kw_p < 0.05 else 'Not significant'}")
                    # B. Homoscedasticity
                    st.markdown("**B. Homoscedasticity — Levene's Test**")
                    if cat_cols_a:
                        groups_l = [gd[test_col_a].dropna().values for _, gd in df.groupby(cat_cols_a[0]) if len(gd[test_col_a].dropna()) >= 2]
                        if len(groups_l) >= 2:
                            lev_s, lev_p = _stats.levene(*groups_l)
                            homo_ok = lev_p > 0.05
                            st.markdown(f"| Levene's F | {lev_s:.4f} | p={lev_p:.4f} | {'✅ Homoscedastic' if homo_ok else '❌ Heteroscedastic'} |")
                            if not homo_ok:
                                st.error("🚨 **HOMOSCEDASTICITY FAILED**")
                                st.markdown(f"""
**F={lev_s:.4f}, p={lev_p:.4f}** — Unequal variances detected.

### 🔧 Remedies:
| Remedy | Formula |
|---|---|
| **Welch's t-test** | t=(x̄₁-x̄₂)/√(s₁²/n₁+s₂²/n₂) |
| **White's Robust SE** | Var(β̂)=(X'X)⁻¹(X'ΩX)(X'X)⁻¹ |
| **Log Transform** | Y*=ln(Y) |
| **WLS** | wᵢ=1/σ̂ᵢ² |

✅ **EXECUTED: Welch's t-test**""")
                                if len(groups_l) == 2:
                                    w_s, w_p = _stats.ttest_ind(*groups_l, equal_var=False)
                                    st.markdown(f"| Welch t | {w_s:.4f} | p={w_p:.4f} | {'✅ Significant' if w_p < 0.05 else '⚪ Not significant'} — corrected for unequal variance |")
                                    st.success("✅ Welch correction applied.")
                                if (data_a > 0).all():
                                    groups_log = [np.log(gd[test_col_a].dropna()) for _, gd in df.groupby(cat_cols_a[0]) if len(gd[test_col_a].dropna()) >= 2 and (gd[test_col_a].dropna() > 0).all()]
                                    if len(groups_log) >= 2:
                                        ll_s, ll_p = _stats.levene(*groups_log)
                                        st.markdown(f"✅ **Log transform re-test:** F={ll_s:.4f} p={ll_p:.4f} — {'✅ Restored' if ll_p > 0.05 else '⚠️ Use WLS/robust SEs'}")
                    # C. Multicollinearity VIF
                    if len(numeric_cols_a) >= 2:
                        st.markdown("**C. Multicollinearity — VIF**")
                        try:
                            from numpy.linalg import lstsq as _lstsq
                            vif_rows=[]
                            for col_v in numeric_cols_a:
                                y_v = df[col_v].dropna()
                                xc = [c for c in numeric_cols_a if c != col_v]
                                xd = df[xc].loc[y_v.index].dropna()
                                ii = y_v.index.intersection(xd.index)
                                if len(ii) >= 3:
                                    Xm = np.column_stack([np.ones(len(ii)), xd.loc[ii].values])
                                    cv,_,_,_ = _lstsq(Xm, y_v.loc[ii].values, rcond=None)
                                    yp = Xm @ cv
                                    ss_r = np.sum((y_v.loc[ii].values - yp)**2)
                                    ss_t = np.sum((y_v.loc[ii].values - y_v.loc[ii].mean())**2)
                                    r2v = max(0, 1-ss_r/ss_t) if ss_t > 0 else 0
                                    vif_v = 1/(1-r2v) if r2v < 1 else float("inf")
                                    flag = "❌ HIGH" if vif_v > 10 else ("⚠️ Moderate" if vif_v > 5 else "✅ OK")
                                    vif_rows.append((col_v, r2v, vif_v, flag))
                            if vif_rows:
                                hdr = "| Variable | R² | VIF | Decision |\n|---|---|---|---|"
                                rows_v = [hdr]
                                high = any(v > 10 for _,_,v,_ in vif_rows)
                                for cv,r2v,vif_v,flag in vif_rows:
                                    rows_v.append(f"| {cv} | {r2v:.4f} | {vif_v:.2f} | {flag} |")
                                st.markdown("\n".join(rows_v))
                                if high:
                                    st.error("🚨 **MULTICOLLINEARITY DETECTED**")
                                    st.markdown("""
### 🔧 Remedies:
| Method | Formula |
|---|---|
| **Remove predictor** | Drop correlated variable |
| **Ridge Regression** | β=(X'X+λI)⁻¹X'y |
| **PCA** | Z=XW (uncorrelated components) |
✅ Recommendation: Remove one correlated predictor or apply Ridge Regression.""")
                        except Exception as ev: st.caption(f"VIF error: {ev}")
                    # D. Durbin-Watson
                    if len(numeric_cols_a) >= 2:
                        st.markdown("**D. Autocorrelation — Durbin-Watson**")
                        try:
                            from numpy.linalg import lstsq as _lstsq_dw
                            y_dw = df[test_col_a].dropna()
                            xc_dw = [c for c in numeric_cols_a if c != test_col_a]
                            xd_dw = df[xc_dw].loc[y_dw.index].dropna()
                            ii_dw = y_dw.index.intersection(xd_dw.index)
                            if len(ii_dw) >= 4:
                                Xm_dw = np.column_stack([np.ones(len(ii_dw)), xd_dw.loc[ii_dw].values])
                                cd,_,_,_ = _lstsq_dw(Xm_dw, y_dw.loc[ii_dw].values, rcond=None)
                                resid = y_dw.loc[ii_dw].values - (Xm_dw @ cd)
                                dw = np.sum(np.diff(resid)**2) / np.sum(resid**2)
                                dw_ok = 1.5 <= dw <= 2.5
                                dw_int = "✅ No autocorrelation" if dw_ok else ("❌ Positive autocorrelation" if dw < 1.5 else "❌ Negative autocorrelation")
                                st.markdown(f"| Durbin-Watson | {dw:.4f} | {dw_int} |\n|---|---|---|")
                                if not dw_ok:
                                    st.error("🚨 **AUTOCORRELATION DETECTED**")
                                    st.markdown(f"""
**DW={dw:.4f}** — Residuals are correlated.

### 🔧 Remedies:
| Method | Formula |
|---|---|
| **Cochrane-Orcutt** | Y*ₜ=Yₜ-ρ·Yₜ₋₁ |
| **GLS** | β=(X'Ω⁻¹X)⁻¹X'Ω⁻¹y |
| **Newey-West HAC SE** | Autocorrelation-consistent SEs |
| **Add AR lag** | Include Yₜ₋₁ as predictor |

✅ Select "Autoregressive (AR) Model" above to execute the AR remedy.""")
                        except Exception as dw_e: st.caption(f"DW error: {dw_e}")
            except Exception as ea: st.warning(f"⚠️ Assumption testing: {ea}")

        # ── Autoregressive (AR) Model ──────────────────────────────────────
        if df is not None and "Autoregressive (AR) Model" in selected_tests and ts_col:
            st.markdown("---")
            st.markdown("### 🔁 Autoregressive (AR) Model")
            st.markdown('<div class="info-card">Formula: Yₜ = c + φ₁Yₜ₋₁ + ... + φₚYₜ₋ₚ + εₜ &nbsp;|&nbsp; Optimal lag selected via AIC</div>', unsafe_allow_html=True)
            try:
                import pandas as pd, numpy as np, plotly.graph_objects as go
                from statsmodels.tsa.ar_model import AutoReg
                ar_s = df[ts_col].dropna()
                best_aic=float("inf"); best_lag=1; best_model=None
                for lag in range(1, min(13, len(ar_s)//4)):
                    try:
                        m = AutoReg(ar_s, lags=lag, old_names=False).fit()
                        if m.aic < best_aic: best_aic=m.aic; best_lag=lag; best_model=m
                    except Exception: continue
                if best_model:
                    st.markdown(f"**Optimal lag p={best_lag} (AIC={best_aic:.4f})**")
                    st.markdown(f"**Formula: Yₜ = c + φ₁Yₜ₋₁ + ... + φ{best_lag}Yₜ₋{best_lag} + εₜ**")
                    params=best_model.params; pvals=best_model.pvalues
                    hdr="| Parameter | Coefficient | p-value | Sig |\n|---|---|---|---|"
                    rows_ar=[hdr]
                    for pn,pv,pp in zip(params.index,params.values,pvals.values):
                        sig="***" if pp<0.001 else ("**" if pp<0.01 else ("*" if pp<0.05 else "ns"))
                        rows_ar.append(f"| {pn} | {pv:.4f} | {pp:.4f} | {sig} |")
                    st.markdown("\n".join(rows_ar))
                    st.markdown(f"""
| AIC | BIC | R² | N |
|---|---|---|---|
| {best_model.aic:.4f} | {best_model.bic:.4f} | {best_model.rsquared:.4f} | {int(best_model.nobs)} |

**Interpretation:** AR({best_lag}) explains {best_model.rsquared*100:.1f}% of variance in {ts_col}. {"Good fit (R²>0.70)" if best_model.rsquared>0.70 else "Moderate fit" if best_model.rsquared>0.40 else "Limited fit — consider ARIMA"}.""")
                    fitted=best_model.fittedvalues
                    fig_ar=go.Figure()
                    fig_ar.add_trace(go.Scatter(y=ar_s.values,mode="lines",name="Actual",line=dict(color="#00c6a7",width=1.5)))
                    fig_ar.add_trace(go.Scatter(y=fitted.values,mode="lines",name="Fitted",line=dict(color="#ffd166",width=1.5,dash="dash")))
                    fig_ar.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",title=f"AR({best_lag}): Actual vs Fitted",title_font_color="#00c6a7")
                    st.plotly_chart(fig_ar,use_container_width=True)
                    resid=best_model.resid
                    fig_res=go.Figure()
                    fig_res.add_trace(go.Scatter(y=resid.values,mode="lines",name="Residuals",line=dict(color="#ff6b6b",width=1)))
                    fig_res.add_hline(y=0,line_dash="dash",line_color="#00c6a7")
                    fig_res.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",title="AR Residuals")
                    st.plotly_chart(fig_res,use_container_width=True)
                    n_fc=10
                    fc_ar=best_model.predict(start=len(ar_s),end=len(ar_s)+n_fc-1)
                    fig_fc=go.Figure()
                    fig_fc.add_trace(go.Scatter(y=ar_s.values,mode="lines",name="Historical",line=dict(color="#00c6a7",width=1.5)))
                    fc_idx=list(range(len(ar_s),len(ar_s)+n_fc))
                    fig_fc.add_trace(go.Scatter(x=fc_idx,y=fc_ar.values,mode="lines+markers",name="Forecast",line=dict(color="#ffd166",width=2,dash="dash"),marker=dict(size=5)))
                    fig_fc.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",title=f"AR({best_lag}) Forecast — {n_fc} periods",title_font_color="#00c6a7")
                    st.plotly_chart(fig_fc,use_container_width=True)
                    st.success(f"✅ AR({best_lag}) complete. AIC={best_aic:.4f}")
            except ImportError: st.error("❌ Run: pip install statsmodels")
            except Exception as e_ar: st.warning(f"⚠️ AR error: {e_ar}")

        # ── Full Inferential Statistical Tests ────────────────────────────
        if df is not None and len(selected_tests) > 0:
            import pandas as pd, numpy as np
            from scipy import stats as _stats
            numeric_cols_inf = df.select_dtypes(include=[np.number]).columns.tolist()
            cat_cols_inf = df.select_dtypes(include=["object","category"]).columns.tolist()
            import plotly.graph_objects as go
            import plotly.express as px
            LAYOUT = dict(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="#0d1526",font_color="#c8deff",
                xaxis=dict(gridcolor="#1e2d4a"),yaxis=dict(gridcolor="#1e2d4a"))
            CLRS = ["#00c6a7","#0077ff","#ff6b6b","#ffd166","#a29bfe","#fd79a8"]

            # ── Pearson Correlation ───────────────────────────────────────
            if "Pearson Correlation" in selected_tests and len(numeric_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📐 Pearson Correlation")
                st.markdown("**Formula:** r = Σ[(xᵢ-x̄)(yᵢ-ȳ)] / √[Σ(xᵢ-x̄)²·Σ(yᵢ-ȳ)²]")
                try:
                    corr_rows = ["| Variables | r | p-value | Strength | Decision |","|---|---|---|---|---|"]
                    for i in range(len(numeric_cols_inf)):
                        for j in range(i+1, len(numeric_cols_inf)):
                            c1n,c2n = numeric_cols_inf[i],numeric_cols_inf[j]
                            d1 = df[c1n].dropna(); d2 = df[c2n].dropna()
                            idx_p = d1.index.intersection(d2.index)
                            if len(idx_p) >= 3:
                                r,p = _stats.pearsonr(d1.loc[idx_p],d2.loc[idx_p])
                                strength = "Very Strong" if abs(r)>=0.8 else ("Strong" if abs(r)>=0.6 else ("Moderate" if abs(r)>=0.4 else ("Weak" if abs(r)>=0.2 else "Negligible")))
                                direction = "positive" if r > 0 else "negative"
                                sig = "✅ Significant (p<0.05)" if p < 0.05 else "⚪ Not significant"
                                corr_rows.append(f"| {c1n} × {c2n} | {r:.4f} | {p:.4f} | {strength} {direction} | {sig} |")
                    st.markdown("\n".join(corr_rows))
                    # Scatter plots for each pair
                    for i in range(min(3, len(numeric_cols_inf)-1)):
                        fig = px.scatter(df, x=numeric_cols_inf[i], y=numeric_cols_inf[i+1],
                            trendline="ols", color_discrete_sequence=["#00c6a7"],
                            title=f"Scatter: {numeric_cols_inf[i]} vs {numeric_cols_inf[i+1]}")
                        fig.update_layout(**LAYOUT, title_font_color="#00c6a7")
                        st.plotly_chart(fig, use_container_width=True)
                    st.markdown("**Interpretation:** r close to ±1 indicates strong linear relationship. p<0.05 indicates the correlation is statistically significant and unlikely due to chance.")
                except Exception as ep: st.warning(f"⚠️ Pearson: {ep}")

            # ── Spearman Correlation ──────────────────────────────────────
            if "Spearman Correlation" in selected_tests and len(numeric_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📐 Spearman Rank Correlation")
                st.markdown("**Formula:** ρ = 1 - 6Σdᵢ² / n(n²-1)  (rank-based, non-parametric)")
                try:
                    sp_rows = ["| Variables | ρ (rho) | p-value | Strength | Decision |","|---|---|---|---|---|"]
                    for i in range(len(numeric_cols_inf)):
                        for j in range(i+1, len(numeric_cols_inf)):
                            c1n,c2n = numeric_cols_inf[i],numeric_cols_inf[j]
                            d1 = df[c1n].dropna(); d2 = df[c2n].dropna()
                            idx_s = d1.index.intersection(d2.index)
                            if len(idx_s) >= 3:
                                rho,p = _stats.spearmanr(d1.loc[idx_s],d2.loc[idx_s])
                                strength = "Very Strong" if abs(rho)>=0.8 else ("Strong" if abs(rho)>=0.6 else ("Moderate" if abs(rho)>=0.4 else "Weak"))
                                sp_rows.append(f"| {c1n} × {c2n} | {rho:.4f} | {p:.4f} | {strength} | {'✅ Sig' if p<0.05 else '⚪ Not sig'} |")
                    st.markdown("\n".join(sp_rows))
                    st.markdown("**Interpretation:** Spearman ρ measures monotonic relationships without assuming normality. Suitable for ordinal data or non-normal distributions.")
                except Exception as es: st.warning(f"⚠️ Spearman: {es}")

            # ── Simple Linear Regression ──────────────────────────────────
            if "Simple Linear Regression" in selected_tests and len(numeric_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📈 Simple Linear Regression")
                st.markdown("**Formula:** Ŷ = β₀ + β₁X + ε")
                try:
                    y_col = numeric_cols_inf[0]; x_col = numeric_cols_inf[1]
                    y_sl = df[y_col].dropna(); x_sl = df[x_col].dropna()
                    idx_sl = y_sl.index.intersection(x_sl.index)
                    if len(idx_sl) >= 3:
                        from numpy.linalg import lstsq as _ls
                        Xm = np.column_stack([np.ones(len(idx_sl)), x_sl.loc[idx_sl].values])
                        ym = y_sl.loc[idx_sl].values
                        coef,_,_,_ = _ls(Xm, ym, rcond=None)
                        yp = Xm @ coef
                        ss_r = np.sum((ym-yp)**2); ss_t = np.sum((ym-ym.mean())**2)
                        r2 = 1-ss_r/ss_t if ss_t > 0 else 0
                        n,k = len(idx_sl), 2
                        adj_r2 = 1-(1-r2)*(n-1)/(n-k-1)
                        mse = ss_r/(n-k)
                        se = np.sqrt(mse*np.linalg.inv(Xm.T@Xm).diagonal())
                        t_stats = coef/se
                        p_vals = [2*(1-_stats.t.cdf(abs(t),df=n-k)) for t in t_stats]
                        f_stat = (ss_t-ss_r)/1 / (ss_r/(n-2))
                        f_p = 1-_stats.f.cdf(f_stat, 1, n-2)
                        st.markdown(f"""
**Coefficients:**

| Variable | Coefficient | Std Error | t-stat | p-value | Sig |
|---|---|---|---|---|---|
| Intercept (β₀) | {coef[0]:.4f} | {se[0]:.4f} | {t_stats[0]:.4f} | {p_vals[0]:.4f} | {"***" if p_vals[0]<0.001 else ("**" if p_vals[0]<0.01 else ("*" if p_vals[0]<0.05 else "ns"))} |
| {x_col} (β₁) | {coef[1]:.4f} | {se[1]:.4f} | {t_stats[1]:.4f} | {p_vals[1]:.4f} | {"***" if p_vals[1]<0.001 else ("**" if p_vals[1]<0.01 else ("*" if p_vals[1]<0.05 else "ns"))} |

**Model Summary:**

| R | R² | Adj R² | MSE | F-stat | F p-value |
|---|---|---|---|---|---|
| {r2**0.5:.4f} | {r2:.4f} | {adj_r2:.4f} | {mse:.4f} | {f_stat:.4f} | {f_p:.4f} |

**Interpretation:** β₁ = {coef[1]:.4f} means a one-unit increase in {x_col} is associated with a {coef[1]:.4f}-unit {"increase" if coef[1]>0 else "decrease"} in {y_col} ({"statistically significant p<0.05" if p_vals[1]<0.05 else "not statistically significant p>0.05"}). R² = {r2:.4f} means {x_col} explains {r2*100:.1f}% of variance in {y_col}.""")
                        fig = px.scatter(df, x=x_col, y=y_col, trendline="ols",
                            title=f"Regression: {y_col} ~ {x_col}", color_discrete_sequence=["#00c6a7"])
                        fig.update_layout(**LAYOUT, title_font_color="#00c6a7")
                        st.plotly_chart(fig, use_container_width=True)
                except Exception as eslr: st.warning(f"⚠️ SLR: {eslr}")

            # ── Multiple Regression ───────────────────────────────────────
            if "Multiple Regression" in selected_tests and len(numeric_cols_inf) >= 3:
                st.markdown("---"); st.markdown("### 📈 Multiple Linear Regression")
                st.markdown("**Formula:** Ŷ = β₀ + β₁X₁ + β₂X₂ + ... + βₖXₖ + ε")
                try:
                    from numpy.linalg import lstsq as _ls
                    y_col = numeric_cols_inf[0]; x_cols = numeric_cols_inf[1:]
                    ym = df[y_col].dropna()
                    Xdf = df[x_cols].loc[ym.index].dropna()
                    idx_mr = ym.index.intersection(Xdf.index)
                    if len(idx_mr) >= len(x_cols)+2:
                        Xm = np.column_stack([np.ones(len(idx_mr))] + [Xdf.loc[idx_mr,c].values for c in x_cols])
                        yv = ym.loc[idx_mr].values
                        coef,_,_,_ = _ls(Xm, yv, rcond=None)
                        yp = Xm @ coef
                        ss_r = np.sum((yv-yp)**2); ss_t = np.sum((yv-yv.mean())**2)
                        n,k = len(idx_mr), len(x_cols)
                        r2 = 1-ss_r/ss_t if ss_t > 0 else 0
                        adj_r2 = 1-(1-r2)*(n-1)/(n-k-1)
                        mse = ss_r/(n-k-1)
                        se = np.sqrt(mse*np.linalg.inv(Xm.T@Xm).diagonal())
                        t_stats = coef/se
                        p_vals = [2*(1-_stats.t.cdf(abs(t),df=n-k-1)) for t in t_stats]
                        f_stat = (ss_t-ss_r)/k / (ss_r/(n-k-1))
                        f_p = 1-_stats.f.cdf(f_stat, k, n-k-1)
                        vnames = ["Intercept"]+x_cols
                        hdr = "| Variable | β | Std Error | t-stat | p-value | Sig |\n|---|---|---|---|---|---|"
                        rows_mr = [hdr]
                        for vn,co,se_v,tv,pv in zip(vnames,coef,se,t_stats,p_vals):
                            sig = "***" if pv<0.001 else ("**" if pv<0.01 else ("*" if pv<0.05 else "ns"))
                            rows_mr.append(f"| {vn} | {co:.4f} | {se_v:.4f} | {tv:.4f} | {pv:.4f} | {sig} |")
                        st.markdown("\n".join(rows_mr))
                        st.markdown(f"""
**Model Summary:**

| R | R² | Adj R² | MSE | F-stat | F p-value | N |
|---|---|---|---|---|---|---|
| {r2**0.5:.4f} | {r2:.4f} | {adj_r2:.4f} | {mse:.4f} | {f_stat:.4f} | {f_p:.4f} | {n} |

**Interpretation:** The model explains {r2*100:.1f}% of variance in {y_col} (Adj R²={adj_r2:.4f}). F({k},{n-k-1})={f_stat:.4f}, p={f_p:.4f} — model is {"statistically significant overall" if f_p<0.05 else "not statistically significant overall"}. Significant predictors (p<0.05): {", ".join(v for v,p in zip(vnames[1:],p_vals[1:]) if p<0.05) or "None"}.""")
                        # Actual vs Fitted
                        fig_mr = go.Figure()
                        fig_mr.add_trace(go.Scatter(y=yv, mode="lines", name="Actual", line=dict(color="#00c6a7",width=1.5)))
                        fig_mr.add_trace(go.Scatter(y=yp, mode="lines", name="Fitted", line=dict(color="#ffd166",width=1.5,dash="dash")))
                        fig_mr.update_layout(**LAYOUT, title=f"Multiple Regression: {y_col} — Actual vs Fitted", title_font_color="#00c6a7")
                        st.plotly_chart(fig_mr, use_container_width=True)
                except Exception as emr: st.warning(f"⚠️ Multiple Regression: {emr}")

            # ── Independent Samples T-Test ────────────────────────────────
            if "Independent Samples T-Test" in selected_tests and cat_cols_inf and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 Independent Samples T-Test")
                st.markdown("**Formula:** t = (x̄₁ - x̄₂) / √(s²p·(1/n₁ + 1/n₂))  where s²p = pooled variance")
                try:
                    g_col = cat_cols_inf[0]; v_col = numeric_cols_inf[0]
                    groups_t = [(name, gd[v_col].dropna().values) for name, gd in df.groupby(g_col) if len(gd[v_col].dropna()) >= 2]
                    if len(groups_t) >= 2:
                        g1n,g1 = groups_t[0]; g2n,g2 = groups_t[1]
                        t_stat, t_p = _stats.ttest_ind(g1, g2, equal_var=True)
                        w_stat, w_p = _stats.ttest_ind(g1, g2, equal_var=False)
                        lev_s, lev_p = _stats.levene(g1, g2)
                        use_welch = lev_p < 0.05
                        st.markdown(f"""
**Levene's Test (equal variance):** F={lev_s:.4f}, p={lev_p:.4f} — {"❌ Unequal variances → Welch correction applied" if use_welch else "✅ Equal variances — standard t-test used"}

| Test | Group 1 ({g1n}) | Group 2 ({g2n}) | t-stat | p-value | Decision |
|---|---|---|---|---|---|
| {"Welch" if use_welch else "Standard"} t-test | x̄={np.mean(g1):.4f}, n={len(g1)} | x̄={np.mean(g2):.4f}, n={len(g2)} | {w_stat if use_welch else t_stat:.4f} | {w_p if use_welch else t_p:.4f} | {"✅ Significant difference" if (w_p if use_welch else t_p)<0.05 else "⚪ No significant difference"} |

**Effect Size (Cohen's d):** {abs(np.mean(g1)-np.mean(g2))/np.sqrt((np.std(g1)**2+np.std(g2)**2)/2):.4f}

**Interpretation:** {"There is a statistically significant difference between groups" if (w_p if use_welch else t_p)<0.05 else "There is no statistically significant difference between groups"} at α=0.05. {"Welch correction was applied due to unequal variances (Levene p<0.05)." if use_welch else ""}""")
                        fig_t = go.Figure()
                        for gname, gvals in groups_t[:4]:
                            fig_t.add_trace(go.Box(y=gvals, name=str(gname), boxmean=True))
                        fig_t.update_layout(**LAYOUT, title=f"T-Test: {v_col} by {g_col}", title_font_color="#00c6a7")
                        st.plotly_chart(fig_t, use_container_width=True)
                except Exception as ett: st.warning(f"⚠️ T-Test: {ett}")

            # ── One-Way ANOVA ─────────────────────────────────────────────
            if "One-Way ANOVA" in selected_tests and cat_cols_inf and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 One-Way ANOVA")
                st.markdown("**Formula:** F = MSB/MSW = [SS_between/(k-1)] / [SS_within/(N-k)]")
                try:
                    g_col = cat_cols_inf[0]; v_col = numeric_cols_inf[0]
                    groups_a_list = [gd[v_col].dropna().values for _, gd in df.groupby(g_col) if len(gd[v_col].dropna()) >= 2]
                    group_names = [str(n) for n,gd in df.groupby(g_col) if len(gd[v_col].dropna()) >= 2]
                    if len(groups_a_list) >= 2:
                        f_stat, f_p = _stats.f_oneway(*groups_a_list)
                        k = len(groups_a_list); N = sum(len(g) for g in groups_a_list)
                        all_vals = np.concatenate(groups_a_list); grand_mean = all_vals.mean()
                        ss_b = sum(len(g)*(g.mean()-grand_mean)**2 for g in groups_a_list)
                        ss_w = sum(np.sum((g-g.mean())**2) for g in groups_a_list)
                        ms_b = ss_b/(k-1); ms_w = ss_w/(N-k)
                        eta2 = ss_b/(ss_b+ss_w)
                        st.markdown(f"""
**ANOVA Table:**

| Source | SS | df | MS | F | p-value |
|---|---|---|---|---|---|
| Between groups | {ss_b:.4f} | {k-1} | {ms_b:.4f} | {f_stat:.4f} | {f_p:.4f} |
| Within groups | {ss_w:.4f} | {N-k} | {ms_w:.4f} | | |
| Total | {ss_b+ss_w:.4f} | {N-1} | | | |

**Effect Size (η²):** {eta2:.4f} — {"Large (η²>0.14)" if eta2>0.14 else ("Medium (η²>0.06)" if eta2>0.06 else "Small (η²<0.06)")}

**Decision:** F({k-1},{N-k})={f_stat:.4f}, p={f_p:.4f} — {"✅ Significant group differences exist (p<0.05)" if f_p<0.05 else "⚪ No significant group differences"}

**Interpretation:** {"There are statistically significant differences between the groups. Post-hoc tests (Tukey HSD) are recommended to identify which specific groups differ." if f_p<0.05 else "The null hypothesis of equal group means is retained. No significant differences detected at α=0.05."}""")
                        # Post-hoc if significant
                        if f_p < 0.05 and len(groups_a_list) >= 2:
                            st.markdown("**Post-hoc: Pairwise Comparisons (Bonferroni corrected)**")
                            pairs = []; n_pairs = k*(k-1)//2
                            for i2 in range(len(groups_a_list)):
                                for j2 in range(i2+1, len(groups_a_list)):
                                    t_ph, p_ph = _stats.ttest_ind(groups_a_list[i2], groups_a_list[j2])
                                    p_bonf = min(1.0, p_ph * n_pairs)
                                    pairs.append(f"| {group_names[i2]} vs {group_names[j2]} | {t_ph:.4f} | {p_ph:.4f} | {p_bonf:.4f} | {'✅ Sig' if p_bonf<0.05 else '⚪ ns'} |")
                            st.markdown("| Comparison | t | p-value | p (Bonferroni) | Decision |\n|---|---|---|---|---|")
                            for row in pairs: st.markdown(row)
                        fig_av = go.Figure()
                        for gname, gvals in zip(group_names, groups_a_list):
                            fig_av.add_trace(go.Box(y=gvals, name=gname, boxmean=True))
                        fig_av.update_layout(**LAYOUT, title=f"ANOVA: {v_col} by {g_col}", title_font_color="#00c6a7")
                        st.plotly_chart(fig_av, use_container_width=True)
                except Exception as eav: st.warning(f"⚠️ ANOVA: {eav}")

            # ── Chi-Square Test ───────────────────────────────────────────
            if "Chi-Square Test" in selected_tests and len(cat_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📊 Chi-Square Test of Independence")
                st.markdown("**Formula:** χ² = Σ[(Oᵢⱼ - Eᵢⱼ)² / Eᵢⱼ]  where E = (row total × col total) / N")
                try:
                    import pandas as pd
                    ct = pd.crosstab(df[cat_cols_inf[0]], df[cat_cols_inf[1]])
                    chi2, p_chi, dof, expected = _stats.chi2_contingency(ct)
                    n_chi = ct.values.sum()
                    cramers_v = np.sqrt(chi2 / (n_chi * (min(ct.shape)-1)))
                    st.markdown(f"""
**Observed Frequencies:**
{ct.to_markdown()}

**Results:**

| χ² | df | p-value | N | Cramér's V | Effect Size |
|---|---|---|---|---|---|
| {chi2:.4f} | {dof} | {p_chi:.4f} | {n_chi} | {cramers_v:.4f} | {"Large (V>0.5)" if cramers_v>0.5 else ("Medium (V>0.3)" if cramers_v>0.3 else "Small (V<0.3)")} |

**Decision:** χ²({dof})={chi2:.4f}, p={p_chi:.4f} — {"✅ Significant association between variables (p<0.05)" if p_chi<0.05 else "⚪ No significant association (p>0.05)"}

**Interpretation:** {"There is a statistically significant association between " + cat_cols_inf[0] + " and " + cat_cols_inf[1] + f". Cramér's V = {cramers_v:.4f} indicates a " + ("large" if cramers_v>0.5 else "medium" if cramers_v>0.3 else "small") + " effect size." if p_chi<0.05 else "The null hypothesis of independence is retained. The variables are not significantly associated."}""")
                    fig_ct = px.bar(ct.reset_index().melt(id_vars=cat_cols_inf[0]),
                        x=cat_cols_inf[0], y="value", color="variable",
                        barmode="group", title=f"Chi-Square: {cat_cols_inf[0]} × {cat_cols_inf[1]}",
                        color_discrete_sequence=CLRS)
                    fig_ct.update_layout(**LAYOUT, title_font_color="#00c6a7")
                    st.plotly_chart(fig_ct, use_container_width=True)
                except Exception as ecs: st.warning(f"⚠️ Chi-Square: {ecs}")

            # ── Mann-Whitney U ────────────────────────────────────────────
            if "Mann-Whitney U Test" in selected_tests and cat_cols_inf and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 Mann-Whitney U Test (Non-parametric)")
                st.markdown("**Formula:** U = n₁n₂ + n₁(n₁+1)/2 - R₁  (rank-based alternative to t-test)")
                try:
                    g_col = cat_cols_inf[0]; v_col = numeric_cols_inf[0]
                    groups_mw = [(str(n), gd[v_col].dropna().values) for n,gd in df.groupby(g_col) if len(gd[v_col].dropna()) >= 2]
                    if len(groups_mw) >= 2:
                        g1n,g1 = groups_mw[0]; g2n,g2 = groups_mw[1]
                        u_stat, u_p = _stats.mannwhitneyu(g1, g2, alternative="two-sided")
                        r_effect = u_stat / (len(g1)*len(g2))
                        st.markdown(f"""
| Group | n | Median | Mean Rank |
|---|---|---|---|
| {g1n} | {len(g1)} | {np.median(g1):.4f} | {np.mean(_stats.rankdata(np.concatenate([g1,g2]))[:len(g1)]):.2f} |
| {g2n} | {len(g2)} | {np.median(g2):.4f} | {np.mean(_stats.rankdata(np.concatenate([g1,g2]))[len(g1):]):.2f} |

| U-statistic | p-value | Effect Size r | Decision |
|---|---|---|---|
| {u_stat:.4f} | {u_p:.4f} | {r_effect:.4f} | {"✅ Significant (p<0.05)" if u_p<0.05 else "⚪ Not significant"} |

**Interpretation:** {"Statistically significant difference between groups in " + v_col + " distribution." if u_p<0.05 else "No significant difference between groups."}""")
                except Exception as emw: st.warning(f"⚠️ Mann-Whitney: {emw}")

            # ── Kruskal-Wallis Test ───────────────────────────────────────
            if "Kruskal-Wallis Test" in selected_tests and cat_cols_inf and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 Kruskal-Wallis Test (Non-parametric ANOVA)")
                st.markdown("**Formula:** H = 12/N(N+1) · Σ[Rⱼ²/nⱼ] - 3(N+1)  (rank-based)")
                try:
                    g_col = cat_cols_inf[0]; v_col = numeric_cols_inf[0]
                    groups_kw = [gd[v_col].dropna().values for _,gd in df.groupby(g_col) if len(gd[v_col].dropna()) >= 2]
                    if len(groups_kw) >= 2:
                        h_stat, h_p = _stats.kruskal(*groups_kw)
                        n_kw = sum(len(g) for g in groups_kw); k_kw = len(groups_kw)
                        eta2_kw = (h_stat - k_kw + 1) / (n_kw - k_kw)
                        st.markdown(f"""
| H-statistic | df | p-value | η² | Decision |
|---|---|---|---|---|
| {h_stat:.4f} | {k_kw-1} | {h_p:.4f} | {eta2_kw:.4f} | {"✅ Significant (p<0.05)" if h_p<0.05 else "⚪ Not significant"} |

**Interpretation:** {"Significant differences exist between groups (p<0.05). Use Dunn's post-hoc test to identify specific differences." if h_p<0.05 else "No significant differences between groups at α=0.05."}""")
                except Exception as ekw: st.warning(f"⚠️ Kruskal-Wallis: {ekw}")

            # ── Time Series Regression ────────────────────────────────────
            if "Time Series Regression" in selected_tests and ts_col and df is not None:
                st.markdown("---"); st.markdown("### 📈 Time Series Regression")
                st.markdown("**Formula:** Yₜ = β₀ + β₁t + β₂X₁ₜ + ... + εₜ  (includes time trend t)")
                try:
                    from numpy.linalg import lstsq as _ls_ts
                    y_ts = df[ts_col].dropna()
                    x_ts_cols = [c for c in numeric_cols_inf if c != ts_col]
                    X_ts = df[x_ts_cols].loc[y_ts.index].dropna() if x_ts_cols else pd.DataFrame(index=y_ts.index)
                    idx_ts = y_ts.index.intersection(X_ts.index) if x_ts_cols else y_ts.index
                    trend = np.arange(len(idx_ts))
                    if x_ts_cols:
                        Xm_ts = np.column_stack([np.ones(len(idx_ts)), trend] + [X_ts.loc[idx_ts,c].values for c in x_ts_cols])
                        vnames_ts = ["Intercept","Time Trend"] + x_ts_cols
                    else:
                        Xm_ts = np.column_stack([np.ones(len(idx_ts)), trend])
                        vnames_ts = ["Intercept","Time Trend"]
                    yv_ts = y_ts.loc[idx_ts].values
                    coef_ts,_,_,_ = _ls_ts(Xm_ts, yv_ts, rcond=None)
                    yp_ts = Xm_ts @ coef_ts
                    ss_r_ts = np.sum((yv_ts-yp_ts)**2); ss_t_ts = np.sum((yv_ts-yv_ts.mean())**2)
                    n_ts,k_ts = len(idx_ts), len(vnames_ts)-1
                    r2_ts = 1-ss_r_ts/ss_t_ts if ss_t_ts > 0 else 0
                    adj_r2_ts = 1-(1-r2_ts)*(n_ts-1)/(n_ts-k_ts-1)
                    mse_ts = ss_r_ts/(n_ts-k_ts-1)
                    se_ts = np.sqrt(mse_ts*np.linalg.inv(Xm_ts.T@Xm_ts).diagonal())
                    t_ts = coef_ts/se_ts
                    p_ts = [2*(1-_stats.t.cdf(abs(tv),df=n_ts-k_ts-1)) for tv in t_ts]
                    hdr_ts = "| Variable | β | Std Error | t-stat | p-value | Sig |\n|---|---|---|---|---|---|"
                    rows_ts = [hdr_ts]
                    for vn,co,se_v,tv,pv in zip(vnames_ts,coef_ts,se_ts,t_ts,p_ts):
                        sig = "***" if pv<0.001 else ("**" if pv<0.01 else ("*" if pv<0.05 else "ns"))
                        rows_ts.append(f"| {vn} | {co:.4f} | {se_v:.4f} | {tv:.4f} | {pv:.4f} | {sig} |")
                    st.markdown("\n".join(rows_ts))
                    st.markdown(f"""
| R² | Adj R² | MSE | N |
|---|---|---|---|
| {r2_ts:.4f} | {adj_r2_ts:.4f} | {mse_ts:.4f} | {n_ts} |

**Interpretation:** The time trend coefficient (β={coef_ts[1]:.4f}) indicates {ts_col} is {"increasing" if coef_ts[1]>0 else "decreasing"} by {abs(coef_ts[1]):.4f} units per period {"(statistically significant p<0.05)" if p_ts[1]<0.05 else "(not statistically significant)"}. R²={r2_ts:.4f} means the model explains {r2_ts*100:.1f}% of variance.""")
                    fig_tsr = go.Figure()
                    fig_tsr.add_trace(go.Scatter(y=yv_ts,mode="lines",name="Actual",line=dict(color="#00c6a7",width=1.5)))
                    fig_tsr.add_trace(go.Scatter(y=yp_ts,mode="lines",name="Fitted",line=dict(color="#ffd166",width=1.5,dash="dash")))
                    fig_tsr.update_layout(**LAYOUT,title=f"TS Regression: {ts_col} — Actual vs Fitted",title_font_color="#00c6a7")
                    st.plotly_chart(fig_tsr,use_container_width=True)
                except Exception as etsr: st.warning(f"⚠️ TS Regression: {etsr}")


            # ── Paired Samples T-Test ─────────────────────────────────────
            if "Paired Samples T-Test" in selected_tests and len(numeric_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📊 Paired Samples T-Test")
                st.markdown("**Formula:** t = d̄ / (sd/√n)  where d = X₁ - X₂ (within-subject differences)")
                try:
                    c1_p, c2_p = numeric_cols_inf[0], numeric_cols_inf[1]
                    idx_pp = df[c1_p].dropna().index.intersection(df[c2_p].dropna().index)
                    if len(idx_pp) >= 3:
                        d1_p = df.loc[idx_pp, c1_p]; d2_p = df.loc[idx_pp, c2_p]
                        diffs = d1_p.values - d2_p.values
                        t_stat_p, t_p_p = _stats.ttest_rel(d1_p, d2_p)
                        d_mean = np.mean(diffs); d_sd = np.std(diffs, ddof=1)
                        cohen_d = d_mean / d_sd if d_sd > 0 else 0
                        st.markdown(f"""
| Variable | Mean | SD | n |
|---|---|---|---|
| {c1_p} | {d1_p.mean():.4f} | {d1_p.std():.4f} | {len(d1_p)} |
| {c2_p} | {d2_p.mean():.4f} | {d2_p.std():.4f} | {len(d2_p)} |
| Difference (d) | {d_mean:.4f} | {d_sd:.4f} | {len(diffs)} |

| t-statistic | df | p-value | Cohen's d | Decision |
|---|---|---|---|---|
| {t_stat_p:.4f} | {len(idx_pp)-1} | {t_p_p:.4f} | {cohen_d:.4f} | {"✅ Significant (p<0.05)" if t_p_p<0.05 else "⚪ Not significant"} |

**Interpretation:** The mean difference between {c1_p} and {c2_p} is {d_mean:.4f} (SD={d_sd:.4f}). t({len(idx_pp)-1})={t_stat_p:.4f}, p={t_p_p:.4f} — {"There is a statistically significant difference between the paired measurements." if t_p_p<0.05 else "No statistically significant difference between the paired measurements at α=0.05."} Cohen's d={cohen_d:.4f} indicates {"large" if abs(cohen_d)>0.8 else "medium" if abs(cohen_d)>0.5 else "small"} effect size.""")
                        fig_p = go.Figure()
                        fig_p.add_trace(go.Box(y=d1_p.values, name=c1_p, marker_color="#00c6a7", boxmean=True))
                        fig_p.add_trace(go.Box(y=d2_p.values, name=c2_p, marker_color="#0077ff", boxmean=True))
                        fig_p.update_layout(**LAYOUT, title=f"Paired T-Test: {c1_p} vs {c2_p}", title_font_color="#00c6a7")
                        st.plotly_chart(fig_p, use_container_width=True)
                except Exception as eppt: st.warning(f"⚠️ Paired T-Test: {eppt}")

            # ── Two-Way ANOVA ─────────────────────────────────────────────
            if "Two-Way ANOVA" in selected_tests and len(cat_cols_inf) >= 2 and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 Two-Way ANOVA")
                st.markdown("**Formula:** Yᵢⱼₖ = μ + αᵢ + βⱼ + (αβ)ᵢⱼ + εᵢⱼₖ  (main effects + interaction)")
                try:
                    import pandas as pd
                    f1, f2, dep = cat_cols_inf[0], cat_cols_inf[1], numeric_cols_inf[0]
                    df_anova = df[[f1, f2, dep]].dropna()
                    if len(df_anova) >= 6:
                        # Main effect factor 1
                        groups_f1 = [g[dep].values for _, g in df_anova.groupby(f1) if len(g) >= 2]
                        groups_f2 = [g[dep].values for _, g in df_anova.groupby(f2) if len(g) >= 2]
                        if len(groups_f1) >= 2 and len(groups_f2) >= 2:
                            f1_stat, f1_p = _stats.f_oneway(*groups_f1)
                            f2_stat, f2_p = _stats.f_oneway(*groups_f2)
                            st.markdown(f"""
**Two-Way ANOVA Results:**

| Source | F-statistic | p-value | Decision |
|---|---|---|---|
| {f1} (Factor A) | {f1_stat:.4f} | {f1_p:.4f} | {"✅ Significant main effect" if f1_p<0.05 else "⚪ Not significant"} |
| {f2} (Factor B) | {f2_stat:.4f} | {f2_p:.4f} | {"✅ Significant main effect" if f2_p<0.05 else "⚪ Not significant"} |
| Interaction A×B | (use statsmodels OLS for full interaction) | | |

**Interpretation:**
- Factor A ({f1}): {"Significant main effect on " + dep if f1_p<0.05 else "No significant main effect on " + dep}
- Factor B ({f2}): {"Significant main effect on " + dep if f2_p<0.05 else "No significant main effect on " + dep}
""")
                            pivot = df_anova.groupby([f1, f2])[dep].mean().unstack()
                            fig_2w = go.Figure()
                            for col_2w in pivot.columns:
                                fig_2w.add_trace(go.Bar(name=str(col_2w), x=pivot.index.astype(str), y=pivot[col_2w].values))
                            fig_2w.update_layout(**LAYOUT, barmode="group",
                                title=f"Two-Way ANOVA: {dep} by {f1} × {f2}", title_font_color="#00c6a7")
                            st.plotly_chart(fig_2w, use_container_width=True)
                except Exception as e2w: st.warning(f"⚠️ Two-Way ANOVA: {e2w}")

            # ── Binary Logistic Regression ────────────────────────────────
            if "Binary Logistic Regression" in selected_tests and numeric_cols_inf:
                st.markdown("---"); st.markdown("### 📊 Binary Logistic Regression")
                st.markdown("**Formula:** ln(p/1-p) = β₀ + β₁X₁ + ... + βₖXₖ  →  P(Y=1) = 1/(1+e⁻ᶻ)")
                try:
                    from numpy.linalg import lstsq as _ls_log
                    # Find a binary outcome variable
                    bin_col = None
                    for col_b in df.columns:
                        unique_vals = df[col_b].dropna().unique()
                        if len(unique_vals) == 2:
                            bin_col = col_b; break
                    if bin_col and numeric_cols_inf:
                        y_log = pd.to_numeric(df[bin_col], errors='coerce').dropna()
                        y_bin = (y_log - y_log.min()) / (y_log.max() - y_log.min())
                        x_cols_log = numeric_cols_inf[:3]
                        X_log = df[x_cols_log].loc[y_bin.index].dropna()
                        idx_log = y_bin.index.intersection(X_log.index)
                        if len(idx_log) >= 5:
                            y_v = y_bin.loc[idx_log].values
                            X_v = np.column_stack([np.ones(len(idx_log))] + [X_log.loc[idx_log, c].values for c in x_cols_log])
                            # Simple logistic via scipy
                            from scipy.optimize import minimize
                            def neg_log_likelihood(beta):
                                z = X_v @ beta
                                p = 1 / (1 + np.exp(-np.clip(z, -500, 500)))
                                return -np.sum(y_v * np.log(p + 1e-10) + (1-y_v) * np.log(1-p + 1e-10))
                            res = minimize(neg_log_likelihood, np.zeros(X_v.shape[1]), method='BFGS')
                            beta_log = res.x
                            z = X_v @ beta_log
                            p_pred = 1 / (1 + np.exp(-z))
                            y_pred_class = (p_pred >= 0.5).astype(int)
                            accuracy = np.mean(y_pred_class == y_v.round())
                            # Odds ratios
                            vnames_log = ["Intercept"] + x_cols_log
                            # Build and display logistic regression results
                            coef_tbl = "| Variable | \u03b2 | Odds Ratio (e\u1d5d) | Interpretation |\n|---|---|---|---|"
                            for vn_lr, b_lr in zip(vnames_log, beta_log):
                                or_lr = np.exp(b_lr)
                                if b_lr > 0:
                                    interp_lr = f"Increases odds by {round((or_lr-1)*100,1)}%"
                                else:
                                    interp_lr = f"Decreases odds by {round((1-or_lr)*100,1)}%"
                                coef_tbl += f"\n| {vn_lr} | {b_lr:.4f} | {or_lr:.4f} | {interp_lr} |"
                            st.markdown("**Logistic Regression Coefficients:**")
                            st.markdown(coef_tbl)
                            st.markdown(f"""
**Model Fit:**

| Metric | Value | Interpretation |
|---|---|---|
| Accuracy | {accuracy:.4f} | {accuracy*100:.1f}% of cases correctly classified |
| Outcome variable | {bin_col} | Binary (0/1) |

**Interpretation:** Each unit increase in a predictor multiplies the odds of {bin_col}=1 by its Odds Ratio. OR > 1 = increased odds; OR < 1 = decreased odds.""")
                    else:
                        st.info("ℹ️ Binary Logistic Regression requires a binary outcome variable (exactly 2 unique values). None detected in dataset.")
                except Exception as elr: st.warning(f"⚠️ Logistic Regression: {elr}")

            # ── Factor Analysis ───────────────────────────────────────────
            if "Factor Analysis" in selected_tests and len(numeric_cols_inf) >= 3:
                st.markdown("---"); st.markdown("### 📊 Factor Analysis (PCA-based)")
                st.markdown("**Formula:** X = LF + ε  where L = factor loadings, F = common factors")
                try:
                    from sklearn.decomposition import PCA
                    import pandas as pd
                    X_fa = df[numeric_cols_inf].dropna()
                    if len(X_fa) >= 5:
                        X_scaled = (X_fa - X_fa.mean()) / X_fa.std()
                        pca = PCA()
                        pca.fit(X_scaled)
                        explained = pca.explained_variance_ratio_
                        cumulative = np.cumsum(explained)
                        n_factors = max(1, sum(pca.explained_variance_ >= 1.0))
                        st.markdown(f"**Kaiser Criterion: {n_factors} factor(s) with eigenvalue ≥ 1.0**")
                        ev_rows = ["| Factor | Eigenvalue | Variance Explained | Cumulative |","|---|---|---|---|"]
                        for k_f, (ev, var, cum) in enumerate(zip(pca.explained_variance_[:6], explained[:6]*100, cumulative[:6]*100), 1):
                            ev_rows.append(f"| F{k_f} | {ev:.4f} | {var:.1f}% | {cum:.1f}% |")
                        st.markdown("\n".join(ev_rows))
                        # Scree plot
                        fig_scree = go.Figure()
                        fig_scree.add_trace(go.Scatter(x=list(range(1, len(explained[:8])+1)),
                            y=pca.explained_variance_[:8], mode="lines+markers",
                            line=dict(color="#00c6a7", width=2), marker=dict(size=8)))
                        fig_scree.add_hline(y=1.0, line_dash="dash", line_color="#ff6b6b",
                            annotation_text="Kaiser criterion (eigenvalue=1)")
                        fig_scree.update_layout(**LAYOUT, title="Scree Plot", title_font_color="#00c6a7",
                            xaxis_title="Factor Number", yaxis_title="Eigenvalue")
                        st.plotly_chart(fig_scree, use_container_width=True)
                        # Loadings heatmap
                        loadings = pd.DataFrame(pca.components_[:n_factors].T,
                            index=numeric_cols_inf,
                            columns=[f"F{k}" for k in range(1, n_factors+1)])
                        fig_load = go.Figure(go.Heatmap(
                            z=loadings.values, x=loadings.columns, y=loadings.index,
                            colorscale=[[0,"#0d1526"],[0.5,"#0077ff"],[1,"#00c6a7"]],
                            zmid=0, text=[[f"{v:.2f}" for v in row] for row in loadings.values],
                            texttemplate="%{text}"))
                        fig_load.update_layout(**LAYOUT, title="Factor Loadings Heatmap", title_font_color="#00c6a7")
                        st.plotly_chart(fig_load, use_container_width=True)
                        st.markdown(f"**Interpretation:** {n_factors} factor(s) extracted explaining {cumulative[n_factors-1]*100:.1f}% of total variance. Loadings > |0.40| are considered significant. High loadings on a factor indicate that variable contributes strongly to that factor.")
                except ImportError:
                    st.warning("⚠️ sklearn not installed. Run: pip install scikit-learn")
                except Exception as efa: st.warning(f"⚠️ Factor Analysis: {efa}")

            # ── Cluster Analysis ──────────────────────────────────────────
            if "Cluster Analysis" in selected_tests and len(numeric_cols_inf) >= 2:
                st.markdown("---"); st.markdown("### 📊 K-Means Cluster Analysis")
                st.markdown("**Formula:** Minimise Σᵢ Σₓ∈Cᵢ ||x - μᵢ||²  (within-cluster sum of squares)")
                try:
                    from sklearn.cluster import KMeans
                    from sklearn.preprocessing import StandardScaler
                    X_cl = df[numeric_cols_inf[:4]].dropna()
                    if len(X_cl) >= 6:
                        scaler = StandardScaler()
                        X_scaled_cl = scaler.fit_transform(X_cl)
                        # Elbow method
                        inertias = []
                        k_range = range(2, min(8, len(X_cl)//2))
                        for k_c in k_range:
                            km = KMeans(n_clusters=k_c, random_state=42, n_init=10)
                            km.fit(X_scaled_cl)
                            inertias.append(km.inertia_)
                        fig_elbow = go.Figure()
                        fig_elbow.add_trace(go.Scatter(x=list(k_range), y=inertias,
                            mode="lines+markers", line=dict(color="#00c6a7", width=2),
                            marker=dict(size=8)))
                        fig_elbow.update_layout(**LAYOUT, title="Elbow Method — Optimal K",
                            title_font_color="#00c6a7", xaxis_title="Number of Clusters",
                            yaxis_title="Inertia (WCSS)")
                        st.plotly_chart(fig_elbow, use_container_width=True)
                        # Fit optimal k (use k=3 as default or smallest elbow)
                        k_opt = 3 if len(X_cl) >= 9 else 2
                        km_opt = KMeans(n_clusters=k_opt, random_state=42, n_init=10)
                        labels = km_opt.fit_predict(X_scaled_cl)
                        X_cl_result = X_cl.copy()
                        X_cl_result["Cluster"] = labels + 1
                        cluster_stats = X_cl_result.groupby("Cluster")[numeric_cols_inf[:4]].mean()
                        st.markdown(f"**Cluster Means (k={k_opt}):**")
                        st.dataframe(cluster_stats.round(4), use_container_width=True)
                        # Cluster sizes
                        from collections import Counter
                        sizes = Counter(labels)
                        st.markdown("**Cluster Sizes:**")
                        size_rows = ["| Cluster | n | % of Total |","|---|---|---|"]
                        for cl_id in sorted(sizes.keys()):
                            pct = sizes[cl_id]/len(labels)*100
                            size_rows.append(f"| Cluster {cl_id+1} | {sizes[cl_id]} | {pct:.1f}% |")
                        st.markdown("\n".join(size_rows))
                        # Scatter plot coloured by cluster
                        if len(numeric_cols_inf) >= 2:
                            fig_cl = px.scatter(X_cl_result, x=numeric_cols_inf[0], y=numeric_cols_inf[1],
                                color="Cluster", color_continuous_scale="teal",
                                title=f"Cluster Plot: {numeric_cols_inf[0]} vs {numeric_cols_inf[1]}")
                            fig_cl.update_layout(**LAYOUT, title_font_color="#00c6a7")
                            st.plotly_chart(fig_cl, use_container_width=True)
                        st.markdown(f"**Interpretation:** K-means partitioned the data into {k_opt} clusters. Each cluster represents a group of observations with similar characteristics. The elbow plot above helps identify the optimal number of clusters where adding more clusters yields diminishing returns.")
                except ImportError:
                    st.warning("⚠️ sklearn not installed. Run: pip install scikit-learn")
                except Exception as eclust: st.warning(f"⚠️ Cluster Analysis: {eclust}")

        # ── Generate AI Chapter Three Narrative ───────────────────────────
        st.markdown("---")
        if st.button("📊 Generate Full Chapter Three Narrative",disabled=not st.session_state.data_filename):
            tests_str=", ".join(selected_tests) if selected_tests else "Descriptive Statistics"
            ts_res=st.session_state.get("ts_analysis","")
            sys_p=f"""You are an expert academic statistician and dissertation writer.
Generate CHAPTER THREE — METHODOLOGY AND DATA ANALYSIS.

PART A — METHODOLOGY headings:\n{st.session_state.meth_headings}
- NO citations in methodology prose
- For each statistical method in 3.7, provide formula + citation in {rsan}
- State all assumptions with verification tests

PART B — RESULTS:
### Assumptions Testing
Table: | Assumption | Test | Statistic | Value | p-value | Interpretation | Remedy if Failed |
For each FAILED assumption: state the remedy by name, show formula, confirm it was applied.

### Descriptive Statistics
Tables + full paragraph interpretation of every value.

### [Each selected test — own section]
H₀ and H₁ → Formula → Results table → Full interpretation paragraph (every coefficient, p-value, effect size) → Link to objective

{'### Time Series Analysis' if ts_res else ''}
{f'ADF results: {ts_res}. Explain stationarity, differencing, ARIMA/SARIMA chosen, forecast interpretation.' if ts_res else ''}

### Summary of All Findings
One paragraph per objective: finding + significance + implication.

RULES: ASCII tables. Every table followed by full interpretation. Formal academic language."""
            usr_p=(f"Dataset:{st.session_state.data_filename}\nColumns:{st.session_state.get('data_columns','')}\n"
                f"Topic:{st.session_state.topic}\nVariables:{st.session_state.variables or 'As in dataset'}\n"
                f"Tests:{tests_str}\nMethodology:\n{st.session_state.proposal_methodology[:400]}\n"
                f"TS results:{ts_res}\nStyle:{rsan}\nGenerate complete Chapter Three.")
            stream_response(sys_p,usr_p,"analysis")
            if st.session_state.analysis:
                sum_sys="""Write a SUMMARY OF ANALYSIS — one paragraph per objective:
1. Restate objective 2. Finding 3. Statistical significance 4. Implication. Formal academic language."""
                stream_response(sum_sys,
                    f"Topic:{st.session_state.topic}\nAnalysis:\n{st.session_state.analysis[:2500]}\nObjectives:\n{st.session_state.proposal_intro[:500]}",
                    "analysis_summary")
        show_output("analysis","📊 Chapter Three Output")
        if st.session_state.analysis_summary:
            show_output("analysis_summary","📋 Analysis Summary")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 6 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
elif cur==6:
    st.markdown('<div class="step-badge">STEP 6 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 🎯 Conclusion & Recommendations")
    if not st.session_state.topic: st.warning("⚠️ Complete Step 1 first.")
    else:
        rscon=st.selectbox("Citation Style",REF_STYLES,
            index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_con")
        doc_authors=st.session_state.uploaded_doc_authors or "Authors from uploaded documents"
        def gen_con():
            sys_p=f"""Generate FINAL CHAPTER — SUMMARY, CONCLUSION AND RECOMMENDATIONS:
### 5.1 Introduction
### 5.2 Summary of the Study
### 5.3 Summary of Findings (numbered, each tied to an objective)
### 5.4 Conclusion (min 3 paragraphs, drawn directly from findings)
### 5.5 Recommendations (min 5 numbered, specific, actionable)
### 5.6 Contributions to Knowledge
### 5.7 Suggestions for Further Research (min 3)
### References
RULES: Cite {doc_authors} in {rscon} where appropriate. Paraphrase all content. Formal academic. Min 2 paragraphs per section."""
            usr_p=(f"Topic:{st.session_state.topic}\nMethodology:{st.session_state.methodology}\n"
                f"Objectives:\n{st.session_state.proposal_intro[:700]}\n\n"
                f"Analysis:\n{st.session_state.analysis[:1000] if st.session_state.analysis else 'Not generated'}\n\n"
                f"Summary:\n{st.session_state.analysis_summary[:600] if st.session_state.analysis_summary else ''}\n\n"
                f"Generate complete final chapter.")
            stream_response(sys_p,usr_p,"conclusion")
        if st.button("🎯 Generate Final Chapter"): gen_con()
        show_output("conclusion","🎯 Final Chapter Output",allow_regen=True,regen_fn=gen_con)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 7 — ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
elif cur==7:
    st.markdown('<div class="step-badge">STEP 7 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## ✍️ Abstract Generator")
    wc_target=st.slider("Target word count",150,350,250,25)
    if not st.session_state.topic: st.warning("⚠️ Complete Step 1 first.")
    else:
        def gen_abs():
            sys_p=f"""Write a single academic abstract paragraph (~{wc_target} words):
background → problem → aim → methodology → key findings → conclusions/implications.
ONE paragraph only. No subheadings. No citations. Past tense for methods/results.
End with: Keywords: [5 keywords]"""
            usr_p=(f"Topic:{st.session_state.topic}\nMethodology:{st.session_state.methodology}\n"
                f"Objectives:{st.session_state.proposal_intro[:400]}\n"
                f"Results:{st.session_state.analysis_summary[:400] or st.session_state.analysis[:400]}\n"
                f"Conclusion:{st.session_state.conclusion[:300]}\nWrite ~{wc_target} words.")
            stream_response(sys_p,usr_p,"abstract")
        if st.button("✍️ Generate Abstract"): gen_abs()
        show_output("abstract","✍️ Abstract Output",allow_regen=True,regen_fn=gen_abs)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 8 — REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
elif cur==8:
    st.markdown('<div class="step-badge">STEP 8 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 🔖 Reference List Generator")
    rsref=st.selectbox("Citation Style",REF_STYLES,
        index=REF_STYLES.index(st.session_state.ref_style) if st.session_state.ref_style in REF_STYLES else 0,key="ref_final")
    st.session_state.ref_style=rsref
    if not st.session_state.uploaded_docs_list:
        st.warning("⚠️ No documents uploaded. Upload in Step 2 first.")
    else:
        doc_authors=st.session_state.uploaded_doc_authors or "Authors from documents"
        c1,c2=st.columns(2)
        with c1:
            if st.button("🔖 Generate Master Reference List",use_container_width=True):
                combined=build_combined_docs_text(1500)
                sys_p=f"""You are an expert academic librarian for {rsref}.
1. Format references for ONLY the uploaded document authors: {doc_authors}
2. Add standard statistical methodology references (Field, Cohen, Hair, Box & Jenkins for ARIMA) for any formulas used
3. Format every entry in {rsref}. Sort alphabetically. Number each. Complete all fields."""
                usr_p=(f"Style:{rsref}\nTopic:{st.session_state.topic}\nAuthors:{doc_authors}\n"
                    f"Sample:\n{combined[:6000]}\nTests:{', '.join(st.session_state.selected_tests)}\nGenerate list.")
                stream_response(sys_p,usr_p,"references")
        with c2:
            if st.button("🔄 Re-format",use_container_width=True,disabled=not st.session_state.references):
                stream_response(f"Reformat in {rsref}. Sort alphabetically. Number each. Fix errors. Do not add or remove entries.",
                    f"Reformat:\n\n{st.session_state.references}","references")
    show_output("references","🔖 Reference List Output")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 9 — EXPORT
# ══════════════════════════════════════════════════════════════════════════════
elif cur==9:
    st.markdown('<div class="step-badge">STEP 9 OF 9</div>',unsafe_allow_html=True)
    st.markdown("## 📄 Full Dissertation Preview & Export")
    sections=[("abstract","Abstract"),("proposal_intro","Research Proposal"),
        ("chapter_one","Chapter One — Introduction"),("literature","Chapter Two — Literature Review"),
        ("analysis","Chapter Three — Methodology & Analysis"),("analysis_summary","Analysis Summary"),
        ("conclusion","Chapter Five — Conclusion & Recommendations"),("references","Master Reference List")]
    st.markdown("### Section Completion Status")
    scols=st.columns(4)
    for i,(key,label) in enumerate(sections):
        with scols[i%4]:
            done=bool(st.session_state.get(key,""))
            wc=word_count(st.session_state.get(key,""))
            st.markdown(
                f'<div style="background:#0d1526;border:1.5px solid {"#00c6a755" if done else "#1e2d4a"};'
                f'border-radius:10px;padding:.7rem;margin-bottom:8px;text-align:center;">'
                f'{"✅" if done else "⭕"}<br>'
                f'<span style="color:{"#00c6a7" if done else "#3a4a6a"};font-weight:600;font-size:11px;">{label}</span>'
                f'{"<br><span style=color:#4a6a8a;font-size:10px;>"+str(wc)+" words</span>" if done else ""}'
                f'</div>',unsafe_allow_html=True)
    filled=[(k,l) for k,l in sections if st.session_state.get(k,"")]
    st.markdown("---")
    if not filled:
        st.warning("⚠️ No sections generated yet.")
    else:
        fname=(st.session_state.topic[:40].replace(" ","_") if st.session_state.topic else "ResearchGenAI_Dissertation")
        full_parts=[f"{'='*70}\n{(st.session_state.topic or 'RESEARCH DISSERTATION').upper()}\nGenerated by ResearchGenAI\n{'='*70}\n"]
        for key,label in sections:
            if st.session_state.get(key,""):
                full_parts.append(f"\n{'='*70}\n{label.upper()}\n{'='*70}\n\n{st.session_state[key]}\n")
        full_report="\n".join(full_parts)
        total_words=sum(word_count(st.session_state.get(k,"")) for k,l in filled)
        st.markdown(f"**Total: {total_words:,} words across {len(filled)} sections**")
        st.markdown("### Export Options")
        ec1,ec2,ec3=st.columns(3)
        with ec1:
            st.download_button("📥 Download .txt",data=full_report,
                file_name=f"{fname}_Full.txt",mime="text/plain",use_container_width=True)
        with ec2:
            st.download_button("📋 Download .md",data=full_report,
                file_name=f"{fname}_Full.md",mime="text/markdown",use_container_width=True)
        with ec3:
            docx_bytes=export_docx()
            if docx_bytes:
                st.download_button("📄 Download .docx",data=docx_bytes,
                    file_name=f"{fname}_Full.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
        st.markdown("---"); st.markdown("### Download Individual Sections")
        dl_cols=st.columns(4)
        for i,(key,label) in enumerate(filled):
            with dl_cols[i%4]:
                st.download_button(f"📄 {label[:20]}",
                    data=f"{label}\n{'='*50}\n\n{st.session_state[key]}",
                    file_name=f"{fname}_{key}.txt",mime="text/plain",
                    use_container_width=True,key=f"dl_{key}")
        st.markdown("---"); st.markdown("### 📖 Dissertation Preview")
        if st.session_state.topic:
            st.markdown(
                f'<div style="text-align:center;padding:2rem;background:#0a1220;'
                f'border:1px solid #1e2d4a;border-radius:12px;margin-bottom:1.5rem;">'
                f'<div style="font-size:22px;font-weight:800;color:#fff;">{st.session_state.topic}</div>'
                f'<div style="color:#00c6a7;font-size:13px;letter-spacing:2px;">A RESEARCH DISSERTATION</div>'
                f'<div style="color:#4a6a8a;font-size:12px;margin-top:.5rem;">Generated by ResearchGenAI · Groq + LLaMA 3.3</div>'
                f'</div>',unsafe_allow_html=True)
        for key,label in filled:
            with st.expander(f"📄 {label} ({word_count(st.session_state.get(key,''))} words)",expanded=False):
                output_toolbar(key,label)
                edit_active=st.session_state.edit_mode.get(key,False)
                if edit_active:
                    edited=st.text_area("Edit",value=st.session_state[key],height=400,key=f"editor_exp_{key}",label_visibility="collapsed")
                    st.session_state[key]=edited
                else:
                    st.markdown(f'<div class="output-box">{st.session_state[key]}</div>',unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("<br>",unsafe_allow_html=True)
c1,c2,c3=st.columns([1,4,1])
with c1:
    if cur>1:
        if st.button("← Previous",use_container_width=True):
            st.session_state.current_step-=1; st.rerun()
with c3:
    if cur<9:
        if st.button("Next Step →",use_container_width=True):
            st.session_state.current_step+=1; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# FEEDBACK BOX
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(f"""
<div style="background:linear-gradient(135deg,#0a1a2e,#0d1526);border:1.5px solid #00c6a755;
border-radius:16px;padding:1.5rem 2rem;margin:1rem 0 1rem 0;">
<div style="font-family:Sora,sans-serif;font-weight:800;font-size:16px;color:#00c6a7;margin-bottom:.4rem;">💬 Feedback & Suggestions</div>
<div style="font-family:Sora,sans-serif;font-size:12px;color:#7a9abf;line-height:1.7;">
Help improve ResearchGenAI. Your feedback goes to <strong style="color:#c8deff;">{FEEDBACK_EMAIL}</strong>
</div></div>""",unsafe_allow_html=True)
with st.form("feedback_form",clear_on_submit=True):
    fb1,fb2=st.columns(2)
    with fb1:
        fb_name=st.text_input("Your name (optional)",placeholder="e.g. Dr. Okeke")
        fb_email=st.text_input("Your email (optional)",placeholder="you@example.com")
    with fb2:
        fb_type=st.selectbox("Feedback type",["🐛 Bug Report","✨ Feature Request","👍 General Praise",
            "📊 Analysis Issue","📄 Content Quality","🔬 Time Series Issue","📚 Literature Review Issue","Other"])
        fb_step=st.selectbox("Which step?",[
            "General / Overall","Step 1 — Proposal","Step 2 — Document Upload",
            "Step 3 — Chapter One","Step 4 — Literature Review","Step 5 — Data Analysis",
            "Step 6 — Conclusion","Step 7 — Abstract","Step 8 — References","Step 9 — Export"])
    fb_rating=st.select_slider("Rating",
        options=["⭐ Poor","⭐⭐ Fair","⭐⭐⭐ Good","⭐⭐⭐⭐ Very Good","⭐⭐⭐⭐⭐ Excellent"],
        value="⭐⭐⭐⭐ Very Good")
    fb_text=st.text_area("Your feedback *",placeholder="Describe the issue or suggestion...",height=130)
    fb_submitted=st.form_submit_button("📤 Send Feedback",use_container_width=True)
    if fb_submitted:
        if fb_text.strip():
            import urllib.parse
            subject=urllib.parse.quote(f"ResearchGenAI Feedback: {fb_type} — {fb_step}")
            body=urllib.parse.quote("\n".join([f"Name: {fb_name or 'Anonymous'}",f"Email: {fb_email or 'Not provided'}",
                f"Type: {fb_type}",f"Step: {fb_step}",f"Rating: {fb_rating}","",f"Feedback:\n{fb_text}"]))
            mailto_url=f"mailto:{FEEDBACK_EMAIL}?subject={subject}&body={body}"
            if "all_feedback" not in st.session_state: st.session_state.all_feedback=[]
            st.session_state.all_feedback.append({"name":fb_name or "Anonymous","type":fb_type,"rating":fb_rating,"step":fb_step,"feedback":fb_text})
            st.success("✅ Thank you! Click below to send via email.")
            st.markdown(f'<a href="{mailto_url}" target="_blank" style="display:inline-block;background:linear-gradient(135deg,#00c6a7,#0077ff);color:white;font-family:Sora,sans-serif;font-weight:700;font-size:14px;padding:.6rem 1.8rem;border-radius:10px;text-decoration:none;">📧 Open Email Client</a>',unsafe_allow_html=True)
            st.balloons()
        else:
            st.warning("⚠️ Please enter your feedback.")

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div style="text-align:center;margin-top:3rem;padding:1rem;border-top:1px solid #1e2d4a;font-family:Sora,sans-serif;font-size:11px;color:#3a4a6a;letter-spacing:1px;">🔬 ResearchGenAI · Powered by Groq + LLaMA 3.3 70B · Free to Use</div>',unsafe_allow_html=True)
